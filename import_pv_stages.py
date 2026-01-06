"""
Script to import Process Validation Stage Templates from Word documents into database
Run this once to populate the pv_stage_templates table
"""

from docx import Document
import sys
import os

# Add parent directory to path to import app modules
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app import app, db
from models import PV_Stage_Template

def import_stages_from_docx(file_path, product_type):
    """Import stage templates from Word document into database"""
    
    print(f"\n{'='*70}")
    print(f"📄 Importing {product_type} stages")
    print(f"📁 File: {file_path}")
    print(f"{'='*70}\n")
    
    try:
        # Read Word document
        doc = Document(file_path)
        
        if len(doc.tables) == 0:
            print(f"❌ No tables found in {file_path}")
            return 0
        
        table = doc.tables[0]
        print(f"✅ Found table with {len(table.rows)} rows\n")
        
        # Skip header row, process data rows
        stages_imported = 0
        
        for i, row in enumerate(table.rows[1:], 1):
            try:
                # Extract data from each column
                stage_name = row.cells[0].text.strip()
                activity = row.cells[1].text.strip()
                key_parameters = row.cells[2].text.strip()
                validation_objective = row.cells[3].text.strip()
                
                # Skip empty rows
                if not stage_name:
                    continue
                
                # Check if stage is optional
                is_optional = "(if applicable)" in stage_name.lower()
                
                # Create stage template object
                stage = PV_Stage_Template(
                    product_type=product_type,
                    stage_number=i,
                    stage_name=stage_name,
                    activity_description=activity,
                    key_parameters=key_parameters,
                    validation_objective=validation_objective,
                    is_optional=is_optional
                )
                
                # Add to database session
                db.session.add(stage)
                stages_imported += 1
                
                # Print progress
                print(f"  ✅ Stage {i:2d}: {stage_name[:60]}")
                
            except Exception as e:
                print(f"  ❌ Error processing row {i}: {e}")
                continue
        
        # Commit all stages for this product type
        db.session.commit()
        print(f"\n✅ Successfully imported {stages_imported} stages for {product_type}")
        
        return stages_imported
        
    except FileNotFoundError:
        print(f"❌ File not found: {file_path}")
        return 0
    except Exception as e:
        print(f"❌ Error reading file: {e}")
        db.session.rollback()
        return 0

def main():
    """Main function to import all stage templates"""
    
    print("\n" + "="*70)
    print("🚀 PROCESS VALIDATION STAGE TEMPLATES IMPORT")
    print("="*70)
    
    with app.app_context():
        # Check if table already has data
        existing_count = PV_Stage_Template.query.count()
        
        if existing_count > 0:
            print(f"\n⚠️  Warning: Found {existing_count} existing stage templates in database")
            response = input("Do you want to delete them and re-import? (yes/no): ")
            
            if response.lower() in ['yes', 'y']:
                print("\n🗑️  Clearing existing stage templates...")
                PV_Stage_Template.query.delete()
                db.session.commit()
                print("✅ Cleared successfully")
            else:
                print("\n❌ Import cancelled. Exiting...")
                return
        
        # Define the 4 stage documents
        stage_docs = {
            'Tablet': r'C:\Users\Admin\OneDrive\Desktop\Process_Validation_Stages_Tablets.docx',
            'Injectable': r'C:\Users\Admin\OneDrive\Desktop\Process_Validation_Stages_Injectables.docx',
            'Capsule': r'C:\Users\Admin\OneDrive\Desktop\Process_Validation_Stages_Capsules.docx',
            'Oral_Liquid': r'C:\Users\Admin\OneDrive\Desktop\Process_Validation_Stages_Oral_Liquids.docx'
        }
        
        total_stages = 0
        successful_imports = 0
        
        # Import each product type
        for product_type, file_path in stage_docs.items():
            try:
                count = import_stages_from_docx(file_path, product_type)
                total_stages += count
                if count > 0:
                    successful_imports += 1
            except Exception as e:
                print(f"\n❌ Critical error importing {product_type}: {e}")
                continue
        
        # Final summary
        print("\n" + "="*70)
        print("📊 IMPORT SUMMARY")
        print("="*70)
        print(f"✅ Product types processed: {successful_imports}/4")
        print(f"✅ Total stages imported: {total_stages}")
        print("="*70)
        
        # Verify import
        print("\n📋 VERIFICATION:")
        print("-"*70)
        for product_type in ['Tablet', 'Injectable', 'Capsule', 'Oral_Liquid']:
            count = PV_Stage_Template.query.filter_by(product_type=product_type).count()
            status = "✅" if count > 0 else "❌"
            print(f"  {status} {product_type:15s}: {count:2d} stages")
        print("-"*70)
        
        print("\n✅ Import process completed!\n")

if __name__ == '__main__':
    main()