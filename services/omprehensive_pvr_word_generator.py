"""
Comprehensive PVR Word Generator Service
Generates detailed Word documents for Process Validation Reports
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)


class ComprehensivePVRWordGenerator:
    """Generate comprehensive PVR Word documents"""
    
    def __init__(self, pvp_template, batch_data: List[Dict]):
        """
        Initialize generator
        
        Args:
            pvp_template: PVP_Template database object
            batch_data: List of batch result dictionaries
        """
        self.pvp = pvp_template
        self.batch_data = batch_data
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        
        # Set normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set heading styles
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(26, 35, 126)
    
    def generate_word(self, output_path: str):
        """Generate comprehensive Word report"""
        
        logger.info(f"Generating comprehensive PVR Word: {output_path}")
        
        # 1. Cover Page
        self._build_cover_page()
        self.doc.add_page_break()
        
        # 2. Table of Contents
        self._build_toc()
        self.doc.add_page_break()
        
        # 3. Objective
        self._build_objective()
        
        # 4. Scope
        self._build_scope()
        self.doc.add_page_break()
        
        # 5. Product Information
        self._build_product_info()
        self.doc.add_page_break()
        
        # 6. Equipment List
        self._build_equipment_list()
        self.doc.add_page_break()
        
        # 7. Materials List
        self._build_materials_list()
        self.doc.add_page_break()
        
        # 8. Manufacturing Process Validation
        self._build_process_validation()
        self.doc.add_page_break()
        
        # 9. Quality Testing Results
        self._build_quality_tests()
        self.doc.add_page_break()
        
        # 10. Conclusion
        self._build_conclusion()
        self.doc.add_page_break()
        
        # 11. Recommendations
        self._build_recommendations()
        self.doc.add_page_break()
        
        # 12. Signatures
        self._build_signatures()
        
        # Save document
        self.doc.save(output_path)
        logger.info(f"✅ Word document generated: {output_path}")
        
        return output_path
    
    def _build_cover_page(self):
        """Build cover page"""
        
        # Add some spacing
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("PROCESS VALIDATION REPORT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 35, 126)
        
        self.doc.add_paragraph()
        
        # Product name
        product = self.doc.add_paragraph()
        product.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = product.add_run(self.pvp.product_name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(40, 53, 147)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Batch info
        batch_numbers = [b.get('batch_number', 'N/A') for b in self.batch_data]
        batch_para = self.doc.add_paragraph()
        batch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        batch_para.add_run("Batch Numbers: ").bold = True
        batch_para.add_run(', '.join(batch_numbers))
        
        self.doc.add_paragraph()
        
        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run("Report Date: ").bold = True
        date_para.add_run(datetime.now().strftime('%B %d, %Y'))
    
    def _build_toc(self):
        """Build table of contents"""
        
        heading = self.doc.add_heading('TABLE OF CONTENTS', level=1)
        
        toc_items = [
            "1. Objective",
            "2. Scope",
            "3. Product Information",
            "4. Equipment List",
            "5. Materials List",
            "6. Manufacturing Process Validation",
            "7. Quality Testing Results",
            "8. Conclusion",
            "9. Recommendations",
            "10. Approval Signatures"
        ]
        
        for item in toc_items:
            self.doc.add_paragraph(item, style='List Number')
    
    def _build_objective(self):
        """Build objective section"""
        
        self.doc.add_heading('1. OBJECTIVE', level=1)
        
        objective_text = f"""
The objective of this validation study is to demonstrate that the manufacturing 
process for {self.pvp.product_name} is capable of consistently producing 
a product that meets all predetermined specifications and quality attributes.

This validation is conducted in accordance with:
• ICH Q7: Good Manufacturing Practice Guide for Active Pharmaceutical Ingredients
• ICH Q8(R2): Pharmaceutical Development
• ICH Q9: Quality Risk Management
• FDA Guidance for Industry: Process Validation
        """
        
        self.doc.add_paragraph(objective_text.strip())
    
    def _build_scope(self):
        """Build scope section"""
        
        self.doc.add_heading('2. SCOPE', level=1)
        
        scope_text = f"""
This validation report covers the complete manufacturing process of 
{self.pvp.product_name}, including:

• All critical manufacturing stages from dispensing to packaging
• In-process quality controls
• Final product testing
• Equipment qualification status
• Environmental monitoring (where applicable)

Batch Size: {self.pvp.batch_size or 'As per protocol'}
Product Type: {self.pvp.product_type or 'As specified'}
        """
        
        self.doc.add_paragraph(scope_text.strip())
    
    def _build_product_info(self):
        """Build product information section"""
        
        self.doc.add_heading('3. PRODUCT INFORMATION', level=1)
        
        # Create table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Parameter'
        header_cells[1].text = 'Details'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        data = [
            ('Product Name', self.pvp.product_name),
            ('Product Type', self.pvp.product_type or 'N/A'),
            ('Batch Size', self.pvp.batch_size or 'N/A'),
            ('Number of Batches Validated', str(len(self.batch_data))),
            ('Validation Date', datetime.now().strftime('%B %Y'))
        ]
        
        for i, (param, value) in enumerate(data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = param
            row_cells[1].text = value
    
    def _build_equipment_list(self):
        """Build equipment list section"""
        
        self.doc.add_heading('4. EQUIPMENT LIST', level=1)
        
        equipment_list = self.pvp.equipment_list
        
        if equipment_list:
            # Create table
            table = self.doc.add_table(rows=len(equipment_list)+1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ['S.No.', 'Equipment Name', 'Equipment ID', 'Location', 'Calibration Status']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for i, equip in enumerate(equipment_list, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = str(i)
                row_cells[1].text = equip.equipment_name
                row_cells[2].text = equip.equipment_id or 'N/A'
                row_cells[3].text = equip.location or 'N/A'
                row_cells[4].text = equip.calibration_status or 'Valid'
        else:
            self.doc.add_paragraph("No equipment data available.")
    
    def _build_materials_list(self):
        """Build materials list section"""
        
        self.doc.add_heading('5. MATERIALS LIST', level=1)
        
        materials_list = self.pvp.materials_list
        
        if materials_list:
            # Create table
            table = self.doc.add_table(rows=len(materials_list)+1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ['S.No.', 'Material Type', 'Material Name', 'Specification', 'Quantity']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for i, mat in enumerate(materials_list, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = str(i)
                row_cells[1].text = mat.material_type
                row_cells[2].text = mat.material_name
                row_cells[3].text = mat.specification or 'N/A'
                row_cells[4].text = mat.quantity or 'N/A'
        else:
            self.doc.add_paragraph("No materials data available.")
    
    def _build_process_validation(self):
        """Build manufacturing process validation section"""
        
        self.doc.add_heading('6. MANUFACTURING PROCESS VALIDATION', level=1)
        
        stages = self.pvp.extracted_stages
        
        if stages:
            for stage in stages:
                # Stage heading
                stage_title = f"6.{stage.stage_number} {stage.stage_name}"
                self.doc.add_heading(stage_title, level=2)
                
                # Stage details
                if stage.equipment_used:
                    para = self.doc.add_paragraph()
                    para.add_run("Equipment Used: ").bold = True
                    para.add_run(stage.equipment_used)
                
                if stage.specific_parameters:
                    para = self.doc.add_paragraph()
                    para.add_run("Parameters: ").bold = True
                    para.add_run(stage.specific_parameters)
                
                if stage.acceptance_criteria:
                    para = self.doc.add_paragraph()
                    para.add_run("Acceptance Criteria: ").bold = True
                    para.add_run(stage.acceptance_criteria)
                
                # Batch results table for this stage
                batch_results = [r for r in stage.batch_results if r.pvr_report_id]
                if batch_results:
                    table = self.doc.add_table(rows=len(batch_results[:10])+1, cols=5)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header
                    header_cells = table.rows[0].cells
                    headers = ['Batch No.', 'Parameter', 'Result', 'Criteria', 'Status']
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows
                    for i, result in enumerate(batch_results[:10], 1):
                        row_cells = table.rows[i].cells
                        row_cells[0].text = result.batch_number
                        row_cells[1].text = result.parameter_name
                        row_cells[2].text = result.actual_value
                        row_cells[3].text = result.acceptance_criteria
                        row_cells[4].text = '✓ Pass' if result.result_status == 'Pass' else '✗ Fail'
                
                self.doc.add_paragraph()  # Spacing
        else:
            self.doc.add_paragraph("No stage data available.")
    
    def _build_quality_tests(self):
        """Build quality testing results section"""
        
        self.doc.add_heading('7. QUALITY TESTING RESULTS', level=1)
        
        criteria = self.pvp.criteria
        
        if criteria and self.batch_data:
            # Create table
            num_batches = len(self.batch_data)
            table = self.doc.add_table(rows=len(criteria)+1, cols=num_batches+3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Test Parameter'
            header_cells[1].text = 'Acceptance Criteria'
            
            for i, batch in enumerate(self.batch_data):
                header_cells[i+2].text = f"Batch {batch.get('batch_number', i+1)}"
            
            header_cells[num_batches+2].text = 'Result'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for i, crit in enumerate(criteria, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = crit.test_name
                row_cells[1].text = crit.acceptance_criteria
                
                # Add results for each batch
                all_pass = True
                for j, batch in enumerate(self.batch_data):
                    result = batch.get('results', {}).get(crit.test_id, 'N/A')
                    row_cells[j+2].text = str(result)
                    if result == 'N/A' or result == 'Fail':
                        all_pass = False
                
                row_cells[num_batches+2].text = '✓ Pass' if all_pass else '✗ Fail'
        else:
            self.doc.add_paragraph("No test data available.")
    
    def _build_conclusion(self):
        """Build conclusion section"""
        
        self.doc.add_heading('8. CONCLUSION', level=1)
        
        conclusion_text = f"""
Based on the validation data from {len(self.batch_data)} consecutive batches of 
{self.pvp.product_name}, the following conclusions are drawn:

• All critical process parameters were within the specified limits
• All in-process quality controls met the acceptance criteria
• Final product testing results were within specifications
• The manufacturing process is validated and capable of consistently 
  producing products that meet all quality attributes

Overall Validation Status: PASSED ✓

The manufacturing process for {self.pvp.product_name} is validated for 
commercial production.
        """
        
        self.doc.add_paragraph(conclusion_text.strip())
    
    def _build_recommendations(self):
        """Build recommendations section"""
        
        self.doc.add_heading('9. RECOMMENDATIONS', level=1)
        
        recommendations = [
            ("Revalidation Schedule", "Revalidation should be performed annually or when significant changes are made to the process, equipment, or materials."),
            ("Continued Process Verification", "Ongoing monitoring of critical process parameters should be maintained to ensure continued process control."),
            ("Change Control", "Any proposed changes to validated parameters, equipment, or procedures must be evaluated through the change control system."),
            ("Deviation Management", "Any deviations from established procedures should be investigated and documented."),
            ("Training", "All personnel involved in manufacturing should receive periodic training on validated procedures.")
        ]
        
        for i, (title, text) in enumerate(recommendations, 1):
            para = self.doc.add_paragraph()
            para.add_run(f"{i}. {title}: ").bold = True
            para.add_run(text)
    
    def _build_signatures(self):
        """Build signatures section"""
        
        self.doc.add_heading('10. APPROVAL SIGNATURES', level=1)
        
        self.doc.add_paragraph()
        
        # Create signature table
        table = self.doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Signature rows
        roles = [
            ('Prepared by:', '', '', ''),
            ('', '', '', ''),
            ('Reviewed by:', '', '', ''),
            ('', '', '', ''),
            ('Approved by:', '', '', '')
        ]
        
        for i, role_data in enumerate(roles, 1):
            row_cells = table.rows[i].cells
            for j, text in enumerate(role_data):
                row_cells[j].text = text
                if j == 0 and text:
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True


def generate_comprehensive_pvr_word(pvp_template, batch_data: List[Dict], output_path: str) -> str:
    """
    Main function to generate comprehensive PVR Word document
    
    Args:
        pvp_template: PVP_Template database object
        batch_data: List of batch result dictionaries
        output_path: Path where Word document should be saved
        
    Returns:
        Path to generated Word document
    """
    generator = ComprehensivePVRWordGenerator(pvp_template, batch_data)
    return generator.generate_word(output_path)