import os
import json
import logging
import pandas as pd
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches
from io import BytesIO
import requests
from services.cloudinary_service import upload_file_from_path, delete_file_by_url
import tempfile
import time
import contextlib
import re
from docx import Document as DocxDocument
from docx.shared import Inches
import PyPDF2
import pdfplumber

@contextlib.contextmanager
def safe_temp_file(suffix='.tmp'):
    """Context manager for safe temporary file handling on Windows."""
    tmp_file = None
    try:
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        yield tmp_file
    finally:
        if tmp_file:
            tmp_file.close()
            # Wait a bit for file handles to be released
            time.sleep(0.1)
            try:
                os.unlink(tmp_file.name)
            except OSError as e:
                logging.warning(f"Could not delete temp file {tmp_file.name}: {e}")
                # Try one more time after a longer delay
                time.sleep(0.5)
                try:
                    os.unlink(tmp_file.name)
                except OSError:
                    logging.warning(f"Failed to delete temp file {tmp_file.name} after retry")

def extract_stp_content(stp_file_url):
    """
    Extract content from STP (Standard Testing Procedure) Word document.
    
    Args:
        stp_file_url: URL of the STP file
    
    Returns:
        dict: Extracted STP content and parameters
    """
    try:
        if not stp_file_url:
            return {}
        
        # Download file content
        response = requests.get(stp_file_url, timeout=30)
        response.raise_for_status()
        
        # Save to temporary file
        with safe_temp_file(suffix='.docx') as tmp_file:
            tmp_file.write(response.content)
            tmp_file.flush()
            
            # Read Word document
            doc = DocxDocument(tmp_file.name)
            
            # Extract text content
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Extract table data
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # Extract key parameters using regex patterns
            text_content = '\n'.join(full_text)
            
            extracted_params = {
                'full_text': text_content,
                'tables': tables_data,
                'parameters': {}
            }
            
            # Common STP parameter patterns
            patterns = {
                'instrument_type': r'(?i)(hplc|gc|uv|titration|ir|aas|lc-ms)',
                'column': r'(?i)column[:\s]+([^\n\r]+)',
                'mobile_phase': r'(?i)mobile\s+phase[:\s]+([^\n\r]+)',
                'flow_rate': r'(?i)flow\s+rate[:\s]+([^\d]+[\d.]+)',
                'detection_wavelength': r'(?i)(?:detection\s+)?wavelength[:\s]+([^\n\r]+)',
                'injection_volume': r'(?i)injection\s+volume[:\s]+([^\n\r]+)',
                'run_time': r'(?i)run\s+time[:\s]+([^\n\r]+)',
                'temperature': r'(?i)temperature[:\s]+([^\n\r]+)',
                'sample_preparation': r'(?i)sample\s+preparation[:\s]+([^\n\r]+)',
                'standard_preparation': r'(?i)standard\s+preparation[:\s]+([^\n\r]+)'
            }
            
            for param_name, pattern in patterns.items():
                match = re.search(pattern, text_content)
                if match:
                    extracted_params['parameters'][param_name] = match.group(1).strip()
            
            logging.info(f"STP content extracted successfully. Found {len(extracted_params['parameters'])} parameters")
            return extracted_params
            
    except Exception as e:
        logging.error(f"Error extracting STP content: {str(e)}")
        return {}

def extract_method_analysis_content(method_analysis_file_url):
    """
    Extract content from Method of Analysis PDF document.
    
    Args:
        method_analysis_file_url: URL of the Method of Analysis PDF file
    
    Returns:
        dict: Extracted method analysis content and parameters
    """
    try:
        if not method_analysis_file_url:
            logging.warning("No method analysis file URL provided")
            return {}
        
        logging.info(f"Starting PDF extraction from URL: {method_analysis_file_url}")
        
        # Check if it's a Cloudinary URL and generate signed URL if needed
        if 'cloudinary.com' in method_analysis_file_url:
            try:
                from services.cloudinary_service import extract_public_id_from_url, generate_signed_url
                
                # Extract public_id from URL
                public_id = extract_public_id_from_url(method_analysis_file_url)
                if public_id:
                    # Determine resource type from URL
                    resource_type = 'raw' if '/raw/' in method_analysis_file_url else 'image'
                    
                    # Generate signed URL
                    signed_url = generate_signed_url(public_id, resource_type=resource_type)
                    if signed_url:
                        logging.info(f"Using signed URL: {signed_url}")
                        method_analysis_file_url = signed_url
                    else:
                        logging.warning("Could not generate signed URL, using original URL")
                else:
                    logging.warning("Could not extract public_id from URL")
            except Exception as e:
                logging.warning(f"Error generating signed URL: {str(e)}, using original URL")
        
        # Download file content
        response = requests.get(method_analysis_file_url, timeout=30)
        response.raise_for_status()
        
        logging.info(f"PDF file downloaded successfully, size: {len(response.content)} bytes")
        
        # Save to temporary file
        with safe_temp_file(suffix='.pdf') as tmp_file:
            tmp_file.write(response.content)
            tmp_file.flush()
            
            # Extract text using pdfplumber (better for tables and formatting)
            extracted_params = {
                'full_text': '',
                'tables': [],
                'parameters': {}
            }
            
            try:
                with pdfplumber.open(tmp_file.name) as pdf:
                    full_text = []
                    tables_data = []
                    
                    logging.info(f"PDF has {len(pdf.pages)} pages")
                    
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            full_text.append(page_text.strip())
                            logging.info(f"Page {page_num + 1}: Extracted {len(page_text)} characters")
                        
                        # Extract tables
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table in page_tables:
                                if table:  # Check if table is not empty
                                    tables_data.append(table)
                                    logging.info(f"Page {page_num + 1}: Found table with {len(table)} rows")
                    
                    extracted_params['full_text'] = '\n'.join(full_text)
                    extracted_params['tables'] = tables_data
                    
                    logging.info(f"Total text extracted: {len(extracted_params['full_text'])} characters")
                    logging.info(f"Total tables found: {len(tables_data)}")
                    
                    # Extract key parameters using regex patterns
                    text_content = extracted_params['full_text']
                    
                    # Enhanced Method of Analysis parameter patterns
                    patterns = {
                        'instrument_type': r'(?i)(hplc|gc|uv|titration|ir|aas|lc-ms|uplc)',
                        'column': r'(?i)column[:\s]+([^\n\r]+)',
                        'mobile_phase': r'(?i)mobile\s+phase[:\s]+([^\n\r]+)',
                        'flow_rate': r'(?i)flow\s+rate[:\s]+([^\d]+[\d.]+)',
                        'detection_wavelength': r'(?i)(?:detection\s+)?wavelength[:\s]+([^\n\r]+)',
                        'injection_volume': r'(?i)injection\s+volume[:\s]+([^\n\r]+)',
                        'run_time': r'(?i)run\s+time[:\s]+([^\n\r]+)',
                        'temperature': r'(?i)temperature[:\s]+([^\n\r]+)',
                        'sample_preparation': r'(?i)sample\s+preparation[:\s]+([^\n\r]+)',
                        'standard_preparation': r'(?i)standard\s+preparation[:\s]+([^\n\r]+)',
                        'method_description': r'(?i)method[:\s]+([^\n\r]+)',
                        'validation_criteria': r'(?i)validation\s+criteria[:\s]+([^\n\r]+)',
                        'acceptance_criteria': r'(?i)acceptance\s+criteria[:\s]*([^\n\r]+)',
                        'procedure': r'(?i)(?:procedure|methodology)[:\s]*([^\n\r]+(?:\n[^\n\r]+)*?)(?=\n\s*(?:acceptance|criteria|end|$))',
                        'titration_procedure': r'(?i)(?:weigh.*?lithium.*?carbonate.*?)([^\n\r]+(?:\n[^\n\r]+)*?)(?=\n\s*(?:acceptance|criteria|end|$))',
                        'equivalence_factor': r'(?i)each\s+ml.*?equivalent\s+to\s+([^\n\r]+)',
                        'indicator': r'(?i)(?:using|with)\s+([^\n\r]*?(?:indicator|solution|methyl\s+orange))'
                    }
                    
                    for param_name, pattern in patterns.items():
                        match = re.search(pattern, text_content)
                        if match:
                            extracted_params['parameters'][param_name] = match.group(1).strip()
                            logging.info(f"Found parameter {param_name}: {match.group(1).strip()}")
                    
                    logging.info(f"Method analysis content extracted successfully. Found {len(extracted_params['parameters'])} parameters")
                    return extracted_params
                    
            except Exception as pdf_error:
                logging.warning(f"pdfplumber failed, trying PyPDF2: {str(pdf_error)}")
                
                # Fallback to PyPDF2
                with open(tmp_file.name, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = []
                    
                    logging.info(f"PyPDF2: PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            full_text.append(page_text.strip())
                            logging.info(f"PyPDF2 Page {page_num + 1}: Extracted {len(page_text)} characters")
                    
                    extracted_params['full_text'] = '\n'.join(full_text)
                    
                    logging.info(f"PyPDF2: Total text extracted: {len(extracted_params['full_text'])} characters")
                    
                    # Extract parameters from text
                    text_content = extracted_params['full_text']
                    patterns = {
                        'instrument_type': r'(?i)(hplc|gc|uv|titration|ir|aas|lc-ms|uplc)',
                        'column': r'(?i)column[:\s]+([^\n\r]+)',
                        'mobile_phase': r'(?i)mobile\s+phase[:\s]+([^\n\r]+)',
                        'flow_rate': r'(?i)flow\s+rate[:\s]+([^\d]+[\d.]+)',
                        'detection_wavelength': r'(?i)(?:detection\s+)?wavelength[:\s]+([^\n\r]+)',
                        'injection_volume': r'(?i)injection\s+volume[:\s]+([^\n\r]+)',
                        'run_time': r'(?i)run\s+time[:\s]+([^\n\r]+)',
                        'temperature': r'(?i)temperature[:\s]+([^\n\r]+)',
                        'sample_preparation': r'(?i)sample\s+preparation[:\s]+([^\n\r]+)',
                        'standard_preparation': r'(?i)standard\s+preparation[:\s]+([^\n\r]+)',
                        'acceptance_criteria': r'(?i)acceptance\s+criteria[:\s]*([^\n\r]+)',
                        'procedure': r'(?i)(?:procedure|methodology)[:\s]*([^\n\r]+(?:\n[^\n\r]+)*?)(?=\n\s*(?:acceptance|criteria|end|$))',
                        'titration_procedure': r'(?i)(?:weigh.*?lithium.*?carbonate.*?)([^\n\r]+(?:\n[^\n\r]+)*?)(?=\n\s*(?:acceptance|criteria|end|$))',
                        'equivalence_factor': r'(?i)each\s+ml.*?equivalent\s+to\s+([^\n\r]+)',
                        'indicator': r'(?i)(?:using|with)\s+([^\n\r]*?(?:indicator|solution|methyl\s+orange))'
                    }
                    
                    for param_name, pattern in patterns.items():
                        match = re.search(pattern, text_content)
                        if match:
                            extracted_params['parameters'][param_name] = match.group(1).strip()
                            logging.info(f"PyPDF2 Found parameter {param_name}: {match.group(1).strip()}")
                    
                    logging.info(f"Method analysis content extracted with PyPDF2. Found {len(extracted_params['parameters'])} parameters")
                    return extracted_params
            
    except Exception as e:
        logging.error(f"Error extracting method analysis content: {str(e)}")
        return {}

def process_raw_data(raw_data_url):
    """
    Process raw data file (CSV/Excel) and extract analytical data.
    
    Args:
        raw_data_url: URL of the raw data file
    
    Returns:
        dict: Processed raw data and statistics
    """
    try:
        if not raw_data_url:
            return {}
        
        # Download file content
        response = requests.get(raw_data_url, timeout=30)
        response.raise_for_status()
        
        # Determine file type and read accordingly
        file_extension = raw_data_url.split('.')[-1].lower()
        
        if file_extension in ['xlsx', 'xls']:
            df = pd.read_excel(BytesIO(response.content))
        elif file_extension == 'csv':
            df = pd.read_csv(BytesIO(response.content))
        else:
            logging.warning(f"Unsupported file type: {file_extension}")
            return {}
        
        if df.empty:
            logging.warning("Raw data file is empty")
            return {}
        
        # Extract basic statistics
        processed_data = {
            'dataframe': df,
            'shape': df.shape,
            'columns': df.columns.tolist(),
            'sample_count': len(df),
            'statistics': {}
        }
        
        # Calculate basic statistics for numeric columns
        numeric_columns = df.select_dtypes(include=['number']).columns
        for col in numeric_columns:
            processed_data['statistics'][col] = {
                'mean': float(df[col].mean()) if not df[col].isna().all() else None,
                'std': float(df[col].std()) if not df[col].isna().all() else None,
                'min': float(df[col].min()) if not df[col].isna().all() else None,
                'max': float(df[col].max()) if not df[col].isna().all() else None,
                'count': int(df[col].count())
            }
        
        # Look for common analytical data patterns
        analytical_data = {}
        
        # Check for peak area data
        peak_area_cols = [col for col in df.columns if 'peak' in col.lower() and 'area' in col.lower()]
        if peak_area_cols:
            analytical_data['peak_areas'] = df[peak_area_cols].to_dict('records')
        
        # Check for retention time data
        rt_cols = [col for col in df.columns if 'retention' in col.lower() or 'rt' in col.lower()]
        if rt_cols:
            analytical_data['retention_times'] = df[rt_cols].to_dict('records')
        
        # Check for concentration data
        conc_cols = [col for col in df.columns if 'concentration' in col.lower() or 'conc' in col.lower()]
        if conc_cols:
            analytical_data['concentrations'] = df[conc_cols].to_dict('records')
        
        processed_data['analytical_data'] = analytical_data
        
        logging.info(f"Raw data processed successfully. Shape: {df.shape}, Columns: {len(df.columns)}")
        return processed_data
        
    except Exception as e:
        logging.error(f"Error processing raw data: {str(e)}")
        return {}

def cleanup_uploaded_files(document):
    """
    Clean up uploaded files after successful document generation.
    This helps save storage space and maintain security.
    
    Args:
        document: Document model instance
    
    Returns:
        dict: Cleanup results
    """
    cleanup_results = {
        'stp_file_deleted': False,
        'raw_data_file_deleted': False,
        'method_analysis_file_deleted': False,
        'errors': []
    }
    
    try:
        # Delete STP file if it exists
        if document.stp_file_url:
            try:
                if delete_file_by_url(document.stp_file_url, resource_type='raw'):
                    cleanup_results['stp_file_deleted'] = True
                    logging.info(f"STP file deleted successfully: {document.stp_file_url}")
                else:
                    cleanup_results['errors'].append("Failed to delete STP file")
            except Exception as e:
                cleanup_results['errors'].append(f"Error deleting STP file: {str(e)}")
        
        # Delete raw data file if it exists
        if document.raw_data_url:
            try:
                if delete_file_by_url(document.raw_data_url, resource_type='raw'):
                    cleanup_results['raw_data_file_deleted'] = True
                    logging.info(f"Raw data file deleted successfully: {document.raw_data_url}")
                else:
                    cleanup_results['errors'].append("Failed to delete raw data file")
            except Exception as e:
                cleanup_results['errors'].append(f"Error deleting raw data file: {str(e)}")
        
        # Delete method analysis file if it exists (for AMV documents)
        if document.method_analysis_file_url:
            try:
                if delete_file_by_url(document.method_analysis_file_url, resource_type='raw'):
                    cleanup_results['method_analysis_file_deleted'] = True
                    logging.info(f"Method analysis file deleted successfully: {document.method_analysis_file_url}")
                else:
                    cleanup_results['errors'].append("Failed to delete method analysis file")
            except Exception as e:
                cleanup_results['errors'].append(f"Error deleting method analysis file: {str(e)}")
        
        # Clear the URLs from database
        document.stp_file_url = None
        document.raw_data_url = None
        document.method_analysis_file_url = None
        
        logging.info(f"Cleanup completed for document {document.id}: {cleanup_results}")
        return cleanup_results
        
    except Exception as e:
        logging.error(f"Error during file cleanup for document {document.id}: {str(e)}")
        cleanup_results['errors'].append(f"General cleanup error: {str(e)}")
        return cleanup_results

def cleanup_old_files(days_old=7):
    """
    Clean up uploaded files for documents that are older than specified days.
    
    Args:
        days_old: Number of days after which to cleanup files (default: 7)
    
    Returns:
        dict: Cleanup statistics
    """
    from database import db
    from models import Document
    from datetime import timedelta
    
    try:
        # Find documents older than specified days that still have uploaded files
        cutoff_date = datetime.now() - timedelta(days=days_old)
        
        old_documents = Document.query.filter(
            Document.created_at < cutoff_date,
            db.or_(
                Document.stp_file_url.isnot(None),
                Document.raw_data_url.isnot(None),
                Document.method_analysis_file_url.isnot(None)
            )
        ).all()
        
        logging.info(f"Found {len(old_documents)} documents with files older than {days_old} days")
        
        cleanup_stats = {
            'total_documents': len(old_documents),
            'successful_cleanups': 0,
            'failed_cleanups': 0,
            'total_files_deleted': 0,
            'errors': []
        }
        
        for document in old_documents:
            try:
                logging.info(f"Cleaning up files for document {document.id}: {document.title}")
                
                cleanup_results = cleanup_uploaded_files(document)
                
                if cleanup_results['errors']:
                    cleanup_stats['failed_cleanups'] += 1
                    cleanup_stats['errors'].extend(cleanup_results['errors'])
                    logging.warning(f"Cleanup failed for document {document.id}: {cleanup_results['errors']}")
                else:
                    cleanup_stats['successful_cleanups'] += 1
                    cleanup_stats['total_files_deleted'] += sum([
                        cleanup_results['stp_file_deleted'],
                        cleanup_results['raw_data_file_deleted'],
                        cleanup_results['method_analysis_file_deleted']
                    ])
                    logging.info(f"Successfully cleaned up document {document.id}")
                
                # Commit changes for each document
                db.session.commit()
                
            except Exception as e:
                cleanup_stats['failed_cleanups'] += 1
                cleanup_stats['errors'].append(f"Error cleaning document {document.id}: {str(e)}")
                logging.error(f"Error cleaning document {document.id}: {str(e)}")
                db.session.rollback()
        
        logging.info(f"Cleanup completed. Stats: {cleanup_stats}")
        return cleanup_stats
        
    except Exception as e:
        logging.error(f"Error during scheduled cleanup: {str(e)}")
        return {'error': str(e)}

def generate_document(document, additional_data=None):
    """
    Generate a pharmaceutical document based on the template and data.

    Args:
        document: Document model instance
        additional_data: Additional data for document generation

    Returns:
        dict: Result with success status and file URLs
    """
    try:
        # Check if Cloudinary is configured
        import os
        if not all([
            os.environ.get('CLOUDINARY_CLOUD_NAME'),
            os.environ.get('CLOUDINARY_API_KEY'),
            os.environ.get('CLOUDINARY_API_SECRET')
        ]):
            return {'success': False, 'error': 'Cloudinary configuration missing. Please check environment variables.'}
        
        if document.document_type == 'AMV':
            return generate_amv_document(document, additional_data)
        elif document.document_type == 'PV':
            return generate_pv_document(document, additional_data)
        elif document.document_type == 'Stability':
            return generate_stability_document(document, additional_data)
        elif document.document_type == 'Degradation':
            return generate_degradation_document(document, additional_data)
        # elif document.document_type == 'Compatibility':  # (when you're ready)
        #     return generate_compatibility_document(document, additional_data)
        else:
            return {'success': False, 'error': 'Unsupported document type'}
    except Exception as e:
        logging.error(f"Document generation error: {str(e)}")
        return {'success': False, 'error': str(e)}


def generate_amv_document(document, additional_data=None):
    """Generate Analytical Method Validation (AMV) document."""
    try:
        # Parse metadata (from create_document_post form fields)
        meta = json.loads(document.document_metadata or "{}")
        title_text = document.title or "Analytical Method Validation Protocol"

        active_ingredient = meta.get("active_ingredient", "")
        strength = meta.get("strength", "")
        pharmacopeia = meta.get("pharmacopeia", "")
        analytical_method = meta.get("analytical_method", "")
        
        # Extract data from uploaded files
        stp_data = {}
        method_analysis_data = {}
        raw_data = {}
        
        # Check for STP file URL in both Document model and AMVDocument model
        stp_file_url = document.stp_file_url
        
        # If not in Document model, check AMVDocument model
        if not stp_file_url:
            from models import AMVDocument
            amv_record = AMVDocument.query.filter_by(document_id=document.id).first()
            if amv_record and amv_record.stp_filename:
                stp_file_url = amv_record.stp_filename
                logging.info(f"Found STP file URL in AMVDocument: {stp_file_url}")
        
        if stp_file_url:
            logging.info("Extracting STP file content...")
            logging.info(f"STP file URL: {stp_file_url}")
            stp_data = extract_stp_content(stp_file_url)
            logging.info(f"STP data extracted: {len(stp_data.get('parameters', {}))} parameters found")
        else:
            logging.warning("No STP file URL found in document")
        
        # Check for method analysis file URL in both Document model and AMVDocument model
        method_analysis_file_url = document.method_analysis_file_url
        
        # If not in Document model, check AMVDocument model
        if not method_analysis_file_url:
            from models import AMVDocument
            amv_record = AMVDocument.query.filter_by(document_id=document.id).first()
            if amv_record and amv_record.method_filename:
                method_analysis_file_url = amv_record.method_filename
                logging.info(f"Found method analysis file URL in AMVDocument: {method_analysis_file_url}")
        
        if method_analysis_file_url:
            logging.info("Extracting Method of Analysis PDF content...")
            logging.info(f"Method analysis file URL: {method_analysis_file_url}")
            method_analysis_data = extract_method_analysis_content(method_analysis_file_url)
            logging.info(f"Method analysis data extracted: {len(method_analysis_data.get('parameters', {}))} parameters found")
            logging.info(f"Method analysis parameters: {method_analysis_data.get('parameters', {})}")
        else:
            # Method PDF is now mandatory - return error if not found
            logging.error("Method analysis PDF is mandatory but not found in document")
            return {"success": False, "error": "Method analysis PDF is required for AMV document generation. Please upload a method analysis PDF."}
        
        if document.raw_data_url:
            logging.info("Processing raw data file...")
            raw_data = process_raw_data(document.raw_data_url)
            logging.info(f"Raw data processed: {raw_data.get('sample_count', 0)} samples found")

        # Create Word document
        doc = DocxDocument()

        # Title
        title = doc.add_heading(title_text.upper(), 0)
        title.alignment = 1  # center

        # Company information with logo
        if document.company:
            # Add company logo if available
            if document.company.logo_url:
                try:
                    response = requests.get(document.company.logo_url, timeout=10)
                    if response.status_code == 200:
                        # Process image to ensure consistent size (150 × 84 pixels)
                        image_data = BytesIO(response.content)
                        
                        # Open image with PIL to resize to exact dimensions
                        from PIL import Image
                        img = Image.open(image_data)
                        
                        # Resize image to exactly 150 × 84 pixels
                        resized_img = img.resize((150, 84), Image.Resampling.LANCZOS)
                        
                        # Save resized image to BytesIO
                        resized_image_data = BytesIO()
                        resized_img.save(resized_image_data, format='PNG')
                        resized_image_data.seek(0)
                        
                        logo_para = doc.add_paragraph()
                        logo_run = logo_para.add_run()
                        logo_run.add_picture(resized_image_data, width=Inches(1.56), height=Inches(0.875))
                        logo_para.alignment = 1  # Center alignment
                except Exception as e:
                    logging.warning(f"Could not load company logo: {e}")
            
            # Add company name and address
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{document.company.name}\n{document.company.address or ''}")
            company_run.bold = True
            company_para.alignment = 1  # Center alignment

        # Product Information
        doc.add_heading("PRODUCT INFORMATION", level=1)
        product_table = doc.add_table(rows=5, cols=2)
        product_table.style = "Table Grid"

        fields = [
            ("Product Name", document.title),
            ("Protocol No.", document.document_number or "AMV/P/001"),
            ("Date", datetime.now().strftime("%d/%m/%Y")),
            ("Active Ingredient", active_ingredient),
            ("Strength", strength),
        ]

        for i, (label, value) in enumerate(fields):
            product_table.rows[i].cells[0].text = label
            product_table.rows[i].cells[1].text = value or "-"

        # Objective
        doc.add_heading("1. OBJECTIVE", level=1)
        doc.add_paragraph(
            "To establish and demonstrate that the validation of the analytical method for the "
            f"{active_ingredient} ({strength}) meets the performance characteristics such as "
            "Specificity, System Suitability, Precision, Linearity, Accuracy, Range, and Robustness."
        )

        # Scope
        doc.add_heading("2. SCOPE", level=1)
        doc.add_paragraph(
            f"This protocol applies to the analytical validation of {active_ingredient} "
            f"using {analytical_method or 'HPLC'} according to {pharmacopeia or 'relevant pharmacopeia'}."
        )

        # Responsibilities
        doc.add_heading("3. RESPONSIBILITY", level=1)
        doc.add_paragraph(
            "QC Analyst:\n"
            "- Prepare protocol and execute validation.\n"
            "- Record all observations.\n\n"
            "Head QC:\n"
            "- Review protocol and approve results.\n\n"
            "QA:\n"
            "- Ensure compliance with SOPs and GLP."
        )

        # Method Parameters (STP + Method of Analysis + Raw Data)
        doc.add_heading("4. METHOD PARAMETERS", level=1)
        
        # Add STP parameters if extracted
        if stp_data and stp_data.get('parameters'):
            doc.add_paragraph("Method parameters extracted from STP document:")
            
            # Create table for STP parameters
            stp_table = doc.add_table(rows=1, cols=2)
            stp_table.style = "Table Grid"
            stp_table.rows[0].cells[0].text = "Parameter"
            stp_table.rows[0].cells[1].text = "Value"
            
            for param_name, param_value in stp_data['parameters'].items():
                row = stp_table.add_row()
                row.cells[0].text = param_name.replace('_', ' ').title()
                row.cells[1].text = param_value
            
            doc.add_paragraph("")  # Add spacing
        
        # Add Method of Analysis parameters if extracted
        if method_analysis_data and method_analysis_data.get('parameters'):
            doc.add_paragraph("Method parameters extracted from Method of Analysis document:")
            
            # Create table for Method of Analysis parameters
            method_table = doc.add_table(rows=1, cols=2)
            method_table.style = "Table Grid"
            method_table.rows[0].cells[0].text = "Parameter"
            method_table.rows[0].cells[1].text = "Value"
            
            for param_name, param_value in method_analysis_data['parameters'].items():
                row = method_table.add_row()
                row.cells[0].text = param_name.replace('_', ' ').title()
                row.cells[1].text = param_value
            
            doc.add_paragraph("")  # Add spacing
            
            # Add specific sections for important parameters
            if 'acceptance_criteria' in method_analysis_data['parameters']:
                doc.add_paragraph("Acceptance Criteria:")
                doc.add_paragraph(method_analysis_data['parameters']['acceptance_criteria'])
                doc.add_paragraph("")
            
            if 'procedure' in method_analysis_data['parameters'] or 'titration_procedure' in method_analysis_data['parameters']:
                doc.add_paragraph("Procedure:")
                procedure_text = method_analysis_data['parameters'].get('procedure') or method_analysis_data['parameters'].get('titration_procedure')
                doc.add_paragraph(procedure_text)
                doc.add_paragraph("")
            
            if 'equivalence_factor' in method_analysis_data['parameters']:
                doc.add_paragraph("Equivalence Factor:")
                doc.add_paragraph(method_analysis_data['parameters']['equivalence_factor'])
                doc.add_paragraph("")
            
            if 'indicator' in method_analysis_data['parameters']:
                doc.add_paragraph("Indicator:")
                doc.add_paragraph(method_analysis_data['parameters']['indicator'])
                doc.add_paragraph("")
        
        # Add raw data summary if processed
        if raw_data and raw_data.get('sample_count', 0) > 0:
            doc.add_paragraph("Analytical data summary:")
            doc.add_paragraph(f"Total samples analyzed: {raw_data['sample_count']}")
            doc.add_paragraph(f"Data columns: {', '.join(raw_data['columns'])}")
            
            # Add statistics table if available
            if raw_data.get('statistics'):
                doc.add_paragraph("Statistical summary:")
                stats_table = doc.add_table(rows=1, cols=6)
                stats_table.style = "Table Grid"
                stats_table.rows[0].cells[0].text = "Parameter"
                stats_table.rows[0].cells[1].text = "Mean"
                stats_table.rows[0].cells[2].text = "Std Dev"
                stats_table.rows[0].cells[3].text = "Min"
                stats_table.rows[0].cells[4].text = "Max"
                stats_table.rows[0].cells[5].text = "Count"
                
                for param, stats in raw_data['statistics'].items():
                    if stats['mean'] is not None:
                        row = stats_table.add_row()
                        row.cells[0].text = param
                        row.cells[1].text = f"{stats['mean']:.3f}" if stats['mean'] else "N/A"
                        row.cells[2].text = f"{stats['std']:.3f}" if stats['std'] else "N/A"
                        row.cells[3].text = f"{stats['min']:.3f}" if stats['min'] else "N/A"
                        row.cells[4].text = f"{stats['max']:.3f}" if stats['max'] else "N/A"
                        row.cells[5].text = str(stats['count'])
        
        # Fallback if no data extracted (should not happen as method PDF is mandatory)
        if not stp_data and not method_analysis_data and not raw_data:
            doc.add_paragraph("ERROR: Method analysis PDF is required but no parameters were extracted. Please ensure a valid method analysis PDF is uploaded.")
            doc.add_paragraph("This document cannot be completed without method parameters from the uploaded PDF.")

        # Validation Parameters Section
        doc.add_heading("5. VALIDATION PARAMETERS", level=1)
        validation_params = meta.get('validation_params', [])
        if validation_params:
            doc.add_paragraph("The following validation parameters will be evaluated:")
            for param in validation_params:
                doc.add_paragraph(f"• {param.replace('_', ' ').title()}")
        else:
            doc.add_paragraph("Validation parameters will be determined based on the analytical method and regulatory requirements.")

        # Parameters to Validate Section
        doc.add_heading("6. PARAMETERS TO VALIDATE/VERIFY", level=1)
        parameters_to_validate = meta.get('parameters_to_validate', [])
        if parameters_to_validate:
            doc.add_paragraph("The following parameters will be validated/verified:")
            for param in parameters_to_validate:
                doc.add_paragraph(f"• {param.replace('_', ' ').title()}")
        else:
            doc.add_paragraph("Parameters to validate will be determined based on the analytical method and regulatory requirements.")

        # Instrument Parameters Section
        doc.add_heading("7. INSTRUMENT PARAMETERS", level=1)
        instrument_params = meta.get('instrument_params', {})
        instrument_type = meta.get('instrument_type', '')
        
        if instrument_params:
            doc.add_paragraph(f"Instrument-specific parameters for {instrument_type.upper()}:")
            
            # Create table for instrument parameters
            inst_table = doc.add_table(rows=1, cols=2)
            inst_table.style = "Table Grid"
            inst_table.rows[0].cells[0].text = "Parameter"
            inst_table.rows[0].cells[1].text = "Value"
            
            for param_name, param_value in instrument_params.items():
                if param_value:  # Only add non-empty values
                    row = inst_table.add_row()
                    row.cells[0].text = param_name.replace('_', ' ').title()
                    row.cells[1].text = str(param_value)
        else:
            doc.add_paragraph("Instrument parameters will be specified based on the selected instrument type and method requirements.")

        # Approval Section
        doc.add_heading("8. APPROVAL", level=1)
        approval_table = doc.add_table(rows=4, cols=4)
        approval_table.style = "Table Grid"

        approval_table.rows[0].cells[0].text = "Name"
        approval_table.rows[0].cells[1].text = "Department"
        approval_table.rows[0].cells[2].text = "Signature"
        approval_table.rows[0].cells[3].text = "Date"

        roles = [
            ("Prepared By", "QC Analyst"),
            ("Reviewed By", "Head QC"),
            ("Approved By", "QA Manager"),
        ]
        for i, (role, dept) in enumerate(roles, start=1):
            row = approval_table.rows[i].cells
            row[0].text = f"{role}\n[Name]"
            row[1].text = dept
            row[2].text = ""
            row[3].text = datetime.now().strftime("%d/%m/%Y")

        # Save Word to temp file and upload
        with safe_temp_file(suffix=".docx") as tmp_file:
            doc.save(tmp_file.name)
            
            # Upload to Cloudinary
            doc_url = upload_file_from_path(
                tmp_file.name,
                folder="generated_documents",
                resource_type="raw"
            )
            
            if not doc_url:
                return {"success": False, "error": "Failed to upload Word document to Cloudinary"}

        # Clean up uploaded files after successful generation
        cleanup_results = cleanup_uploaded_files(document)
        if cleanup_results['errors']:
            logging.warning(f"Some files could not be cleaned up: {cleanup_results['errors']}")
        else:
            logging.info("All uploaded files cleaned up successfully")

        return {
            "success": True, 
            "doc_url": doc_url, 
            "excel_url": None,  # No Excel file generated
            "cleanup_results": cleanup_results
        }

    except Exception as e:
        logging.error(f"AMV document generation error: {str(e)}")
        return {"success": False, "error": str(e)}

def generate_amv_excel(document, additional_data=None):
    """Generate Excel file with AMV calculations (with metadata, formulas, formatting)."""
    try:
        # Extract data from uploaded files for Excel
        stp_data = {}
        raw_data = {}
        
        if document.stp_file_url:
            stp_data = extract_stp_content(document.stp_file_url)
        
        if document.raw_data_url:
            raw_data = process_raw_data(document.raw_data_url)
        
        with safe_temp_file(suffix='.xlsx') as tmp_file:
            writer = pd.ExcelWriter(tmp_file.name, engine='openpyxl')

            # --------------------
            # 1. Protocol Info
            # --------------------
            info_data = {
                'Field': ['Title', 'Document No.', 'Product', 'Active Ingredient', 'Company', 'Date'],
                'Value': [
                    document.title,
                    document.document_number or '',
                    (json.loads(document.document_metadata).get('product_name', '') if document.document_metadata else ''),
                    (json.loads(document.document_metadata).get('active_ingredient', '') if document.document_metadata else ''),
                    document.company.name if document.company else '',
                    datetime.now().strftime('%d-%m-%Y')
                ]
            }
            pd.DataFrame(info_data).to_excel(writer, sheet_name='Protocol Info', index=False)

            # --------------------
            # 2. System Suitability (with actual data if available)
            # --------------------
            ss_data = {
                'Parameter': ['Retention Time', 'Peak Area', 'Tailing Factor', 'Theoretical Plates'],
                'Injection 1': ['', '', '', ''],
                'Injection 2': ['', '', '', ''],
                'Injection 3': ['', '', '', ''],
                'Injection 4': ['', '', '', ''],
                'Injection 5': ['', '', '', ''],
                'Mean': ['', '', '', ''],
                '%RSD': ['', '', '', ''],
                'Acceptance Criteria': ['RSD ≤ 2%', 'RSD ≤ 2%', 'NMT 2.0', 'NLT 2000']
            }
            
            # Populate with actual data if available
            if raw_data and raw_data.get('analytical_data'):
                analytical_data = raw_data['analytical_data']
                
                # Fill retention times if available
                if 'retention_times' in analytical_data:
                    rt_data = analytical_data['retention_times']
                    for i, rt_record in enumerate(rt_data[:5]):  # First 5 injections
                        for col_name, value in rt_record.items():
                            if 'retention' in col_name.lower() or 'rt' in col_name.lower():
                                ss_data[f'Injection {i+1}'][0] = f"{value:.3f}" if value else ""
                
                # Fill peak areas if available
                if 'peak_areas' in analytical_data:
                    pa_data = analytical_data['peak_areas']
                    for i, pa_record in enumerate(pa_data[:5]):  # First 5 injections
                        for col_name, value in pa_record.items():
                            if 'peak' in col_name.lower() and 'area' in col_name.lower():
                                ss_data[f'Injection {i+1}'][1] = f"{value:.0f}" if value else ""
            
            pd.DataFrame(ss_data).to_excel(writer, sheet_name='System Suitability', index=False)

            # --------------------
            # 2.5. STP Parameters (if extracted)
            # --------------------
            if stp_data and stp_data.get('parameters'):
                stp_params_data = {
                    'Parameter': list(stp_data['parameters'].keys()),
                    'Value': list(stp_data['parameters'].values())
                }
                pd.DataFrame(stp_params_data).to_excel(writer, sheet_name='STP Parameters', index=False)

            # --------------------
            # 3. Precision
            # --------------------
            precision_data = {
                'Sample': [f'Sample {i+1}' for i in range(6)],
                'Peak Area': [''] * 6,
                'Assay %': [''] * 6
            }
            pd.DataFrame(precision_data).to_excel(writer, sheet_name='Precision', index=False)

            # --------------------
            # 4. Linearity
            # --------------------
            linearity_data = {
                'Concentration (%)': [50, 75, 100, 125, 150],
                'Peak Area': [''] * 5,
                'Response Factor': [''] * 5
            }
            pd.DataFrame(linearity_data).to_excel(writer, sheet_name='Linearity', index=False)

            # --------------------
            # 5. Accuracy
            # --------------------
            accuracy_data = {
                'Level (%)': [50, 100, 150],
                'Amount Added': [''] * 3,
                'Amount Found': [''] * 3,
                'Recovery %': [''] * 3
            }
            pd.DataFrame(accuracy_data).to_excel(writer, sheet_name='Accuracy', index=False)

            # --------------------
            # 6. Robustness
            # --------------------
            robustness_data = {
                'Parameter': ['Flow Rate +10%', 'Flow Rate -10%', 'Wavelength +2nm', 'Wavelength -2nm'],
                'Retention Time': [''] * 4,
                'Peak Area': [''] * 4,
                'Acceptance Criteria': ['No significant change'] * 4
            }
            pd.DataFrame(robustness_data).to_excel(writer, sheet_name='Robustness', index=False)

            # --------------------
            # 7. LOD & LOQ
            # --------------------
            lod_loq_data = {
                'Parameter': ['LOD', 'LOQ'],
                'Concentration': ['', ''],
                'Peak Area': ['', ''],
                'Acceptance Criteria': ['S/N ≥ 3', 'S/N ≥ 10']
            }
            pd.DataFrame(lod_loq_data).to_excel(writer, sheet_name='LOD_LOQ', index=False)

            writer.close()

            # --------------------
            # Add formulas + styling with openpyxl
            # --------------------
            from openpyxl import load_workbook
            from openpyxl.styles import Font, Alignment

            wb = load_workbook(tmp_file.name)

            # Format headers + add formulas for System Suitability
            ws = wb['System Suitability']
            for cell in ws[1]:
                cell.font = Font(bold=True)
                ws.column_dimensions[cell.column_letter].width = 18

            # Insert Mean & %RSD formulas dynamically (rows 2–5)
            for row in range(2, 6):
                ws[f'H{row}'] = f"=AVERAGE(B{row}:F{row})"
                ws[f'I{row}'] = f"=STDEV(B{row}:F{row})/H{row}*100"
                ws[f'H{row}'].font = Font(bold=True)
                ws[f'I{row}'].font = Font(bold=True)

            # Save workbook
            wb.save(tmp_file.name)

            # --------------------
            # Upload to Cloudinary
            # --------------------
            excel_url = upload_file_from_path(tmp_file.name, folder='generated_excel', resource_type='raw')
            return excel_url

    except Exception as e:
        logging.error(f"Excel generation error: {str(e)}")
        return None

def generate_pv_document(document, additional_data=None):
    """Generate Process Validation document."""
    # Similar structure to AMV but with PV-specific content
    return {'success': False, 'error': 'PV document generation not yet implemented'}

def generate_stability_document(document, additional_data=None):
    """Generate Stability Study document."""
    # Similar structure to AMV but with stability-specific content
    return {'success': False, 'error': 'Stability document generation not yet implemented'}

def generate_degradation_document(document, additional_data=None):
    """Generate Forced Degradation document."""
    # Similar structure to AMV but with degradation-specific content
    return {'success': False, 'error': 'Degradation document generation not yet implemented'}

def download_file_content(url):
    """Download file content from URL."""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except Exception as e:
        logging.error(f"Error downloading file: {str(e)}")
        return None

def process_raw_data_for_amv(doc, raw_data_content):
    """Process raw data and add to AMV document."""
    try:
        # Try to read as Excel first, then CSV
        try:
            df = pd.read_excel(BytesIO(raw_data_content))
        except:
            df = pd.read_csv(BytesIO(raw_data_content))

        if not df.empty:
            doc.add_paragraph("Raw data has been processed and incorporated into the method validation calculations.")

            # Add a sample of the data
            if len(df) > 0:
                doc.add_paragraph(f"Data contains {len(df)} rows and {len(df.columns)} columns.")
                doc.add_paragraph(f"Columns: {', '.join(df.columns.tolist())}")

    except Exception as e:
        logging.error(f"Error processing raw data: {str(e)}")
        doc.add_paragraph("Error processing raw data. Please check data format.")
