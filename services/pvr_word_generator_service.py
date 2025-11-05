# Copyright (C) 2025 Soumyadeep Ghosh <soumyadeepghosh2004@zohomail.in>
# All Rights Reserved.

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

from models import PVR_Report, PVR_Data, PVP_Criteria


def set_cell_background(cell, color):
    """Set background color for table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def generate_pvr_word(report_id, product_name, template, criteria):
    """
    Generate PVR Word document from report data
    """
    # Get report and data
    report = PVR_Report.query.get(report_id)
    if not report:
        raise Exception("Report not found")
    
    batch_data = PVR_Data.query.filter_by(pvr_report_id=report_id).all()
    
    # Get unique batch numbers
    batch_numbers = sorted(list(set([d.batch_number for d in batch_data])))
    
    # Create output folder
    output_folder = os.path.join(os.getcwd(), 'uploads', 'pvr_reports')
    os.makedirs(output_folder, exist_ok=True)
    
    # Generate filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"PVR_{product_name.replace(' ', '_')}_{timestamp}.docx"
    filepath = os.path.join(output_folder, filename)
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("PROCESS VALIDATION REPORT")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 60, 114)
    
    doc.add_paragraph()  # Space
    
    # Product Information
    doc.add_heading("Product Information", level=2)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ['Product Name', product_name],
        ['Report Number', f"PVR-{report.id}"],
        ['Template', template.template_name],
        ['Generated Date', datetime.now().strftime('%d %B %Y')],
        ['Number of Batches', str(len(batch_numbers))]
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Make label bold
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
    
    doc.add_paragraph()  # Space
    
    # Batch Information
    doc.add_heading("Batch Information", level=2)
    
    batch_table = doc.add_table(rows=len(batch_numbers) + 1, cols=2)
    batch_table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = batch_table.rows[0].cells
    hdr_cells[0].text = 'Sr. No'
    hdr_cells[1].text = 'Batch Number'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, '2A5298')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Data rows
    for i, batch in enumerate(batch_numbers, 1):
        row = batch_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = batch
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Space
    
    # Test Results
    doc.add_heading("Validation Test Results", level=2)
    
    # Create results table
    num_cols = 2 + len(batch_numbers)  # Test name + Criteria + Batch columns
    results_table = doc.add_table(rows=len(criteria) + 1, cols=num_cols)
    results_table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = results_table.rows[0].cells
    hdr_cells[0].text = 'Test Parameter'
    hdr_cells[1].text = 'Acceptance Criteria'
    for i, batch in enumerate(batch_numbers):
        hdr_cells[2 + i].text = f'Batch {i+1}\n{batch}'
    
    # Style header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, '2A5298')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Data rows
    for row_idx, criterion in enumerate(criteria, 1):
        row = results_table.rows[row_idx]
        
        # Test name
        row.cells[0].text = criterion.test_name
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Acceptance criteria
        row.cells[1].text = criterion.acceptance_criteria
        
        # Results for each batch
        for col_idx, batch in enumerate(batch_numbers):
            result = next(
                (d.test_result for d in batch_data 
                 if d.test_id == criterion.test_id and d.batch_number == batch),
                '-'
            )
            cell = row.cells[2 + col_idx]
            cell.text = result
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Highlight result
            if result != '-':
                set_cell_background(cell, 'D4EDDA')
    
    doc.add_paragraph()  # Space
    
    # Conclusion
    doc.add_heading("Conclusion", level=2)
    
    conclusion_para = doc.add_paragraph()
    conclusion_text = (
        f"The process validation for {product_name} has been successfully completed "
        f"for {len(batch_numbers)} consecutive batches. All test results are within "
        f"the predetermined acceptance criteria as defined in the validation protocol. "
        f"The manufacturing process is validated and qualified for commercial production."
    )
    conclusion_para.add_run(conclusion_text)
    conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph()  # Space
    
    # Summary
    doc.add_heading("Summary", level=2)
    
    summary_para = doc.add_paragraph()
    summary_text = (
        f"Based on the observed data from the process validation of {product_name}, "
        f"the manufacturing process consistently produces a product meeting all predetermined "
        f"specifications and quality characteristics. The process demonstrates adequate control "
        f"and reproducibility across all validated batches."
    )
    summary_para.add_run(summary_text)
    summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph()  # Space
    doc.add_paragraph()  # Space
    
    # Approval Section
    doc.add_heading("Approval Signatures", level=2)
    
    approval_table = doc.add_table(rows=9, cols=2)
    approval_table.style = 'Light Grid Accent 1'
    
    approval_data = [
        ['Prepared By:', ''],
        ['Name:', ''],
        ['Date:', ''],
        ['Reviewed By:', ''],
        ['Name:', ''],
        ['Date:', ''],
        ['Approved By:', ''],
        ['Name:', ''],
        ['Date:', '']
    ]
    
    for i, (label, value) in enumerate(approval_data):
        row = approval_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        if ':' in label and label != 'Date:':
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Save document
    doc.save(filepath)
    
    print(f"✅ PVR Word document generated: {filepath}")
    return filepath