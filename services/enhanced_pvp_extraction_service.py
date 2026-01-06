"""
Enhanced PVP AI Extraction Service
Extracts comprehensive data from Process Validation Protocol PDFs using AI
"""

import os
import re
import json
import math
import logging
import datetime
from typing import Dict, List, Optional
from pathlib import Path

import pdfplumber
import pandas as pd

# Optional AI + OCR + PDF rendering
try:
    import camelot
except Exception:
    camelot = None

try:
    import google.generativeai as genai
except Exception:
    genai = None

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception:
    Document = None

try:
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML
except Exception:
    Environment = None
    HTML = None

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
logging.getLogger('pdfminer').setLevel(logging.ERROR)
logging.getLogger('pdfplumber').setLevel(logging.ERROR)

# Configure Gemini (optional)
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
if GEMINI_API_KEY and genai:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-pro')
    except Exception as e:
        logger.warning("Failed to configure Gemini model: %s", e)
        model = None
else:
    model = None
    logger.info("Gemini API key not found or SDK not available. Using regex-only extraction.")


# -----------------------
# Utility helpers
# -----------------------
def _to_float(x):
    try:
        s = str(x)
        s = s.replace(',', '')
        v = re.search(r'-?\d+(\.\d+)?', s)
        return float(v.group(0)) if v else None
    except Exception:
        return None


def compute_stats(values):
    vals = [v for v in (_to_float(x) for x in values) if v is not None]
    if not vals:
        return {}
    mean = sum(vals) / len(vals)
    var = sum((v - mean) ** 2 for v in vals) / len(vals)
    std = math.sqrt(var)
    rsd = (std / mean * 100) if mean != 0 else None
    return {
        'mean': round(mean, 6),
        'std': round(std, 6),
        'rsd': round(rsd, 4) if rsd is not None else None,
        'count': len(vals)
    }


def normalize_header(h: str) -> str:
    return re.sub(r'[^a-z0-9 ]', '', str(h).lower()).strip()


def ensure_df(table):
    """Return a pandas DataFrame for a camelot table or None if empty."""
    if table is None:
        return None
    df = getattr(table, 'df', None)
    if df is None:
        return None
    try:
        df = pd.DataFrame(df)
    except Exception:
        pass
    if df is None or df.shape[0] == 0:
        return None
    return df


def _clean_text_val(v: Optional[str], max_len: int = 190) -> str:
    """Normalize whitespace, remove control chars, truncate to max_len."""
    if v is None:
        return ''
    s = str(v)
    s = s.replace('\r', ' ').replace('\n', ' ').strip()
    s = re.sub(r'\s+', ' ', s)
    s = ''.join(ch for ch in s if ch.isprintable())
    s = s.strip()
    if len(s) > max_len:
        s = s[:max_len].rstrip()
    return s


def _is_obvious_heading(s: str) -> bool:
    if not s:
        return False
    low = s.lower()
    skip_prefixes = (
        'cover page', 'table of contents', 'protocol approval', 'objective',
        'scope', 'validation approach', 'reason for validation', 'revalidation',
        'index', 'page', 'product standard', 'stability protocol', 'contents',
        'cover', 'acknowledgement', 'references', 'appendix', 'format no',
        'effective date', 'supersedes'
    )
    for p in skip_prefixes:
        if low.startswith(p):
            return True
    if re.fullmatch(r'[0-9\.\-]{1,4}', low):
        return True
    return False


# -----------------------
# Main extractor
# -----------------------
class EnhancedPVPExtractor:
    """Extract comprehensive data from PVP documents"""

    def __init__(self, pdf_path: str, tesseract_cmd: Optional[str] = None):
        self.pdf_path = str(pdf_path)
        self.full_text: str = ""
        self.product_type: Optional[str] = None
        self.tables = []
        self.tables_df: List[pd.DataFrame] = []
        if tesseract_cmd:
            os.environ['TESSERACT_CMD'] = tesseract_cmd
            if pytesseract:
                try:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                except Exception:
                    pass

    # -----------------------
    # Public runner
    # -----------------------
    def extract_all(self) -> Dict:
        """Main extraction method - extracts everything from PVP"""
        logger.info(f"Starting extraction from: {self.pdf_path}")

        # Extract full text from PDF (with OCR fallback)
        self.full_text = self._extract_text_from_pdf()
        if not self.full_text or len(self.full_text.strip()) == 0:
            logger.error("Failed to extract text from PDF or it's empty")
            return self._empty_result()
        logger.info(f"Extracted text length: {len(self.full_text)}")

        # Extract tables
        self.tables = self._extract_tables_from_pdf()
        self.tables_df = []
        for t in self.tables:
            df = ensure_df(t)
            if df is not None:
                self.tables_df.append(df)

        logger.info(f"Extracted {len(self.tables_df)} usable tables from PDF")

        # Build result with ALL sections
        result = {
            'product_info': self._extract_product_info(),
            'product_type': self._detect_product_type(),
            'company_info': self._extract_company_info(),
            'equipment': self._extract_equipment(),
            'materials': self._extract_materials(),
            'stages': self._extract_stages(),
            'test_criteria': self._extract_test_criteria(),
            'test_preparations': self._extract_test_preparations(),
            'calculation_sheet': self._extract_calculation_sheet(),
            'batch_details': self._extract_batch_details(),

            # NEW SECTIONS
            'hold_time_study': self._extract_hold_time_study(),
            'bioburden_bet': self._extract_bioburden_bet(),
            'qc_final_product': self._extract_qc_final_product(),
            'water_quality': self._extract_water_quality(),
            'vial_sterility': self._extract_vial_sterility(),
            'manufacturing_params': self._extract_manufacturing_params(),

            'observations': self._extract_observations(),
            'signatures': self._extract_signatures(),
            'protocol_summary': self._extract_protocol_summary(),
            'raw_text_length': len(self.full_text)
        }

        # Calculate statistics
        result['statistics'] = self._calculate_statistics_from_criteria(result['test_criteria'])

        self.product_type = result['product_type']
        logger.info("Extraction complete: product=%s, type=%s",
                    result['product_info'].get('product_name'), result['product_type'])
        return result

    # -----------------------
    # Text extraction with OCR fallback
    # -----------------------
    def _extract_text_from_pdf(self) -> str:
        text = ""
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:
                        page_text = ""

                    # If text is short or empty, try OCR
                    if not page_text or len(page_text.strip()) < 60:
                        if pytesseract:
                            try:
                                pil_img = page.to_image(resolution=200).original
                                ocr_text = pytesseract.image_to_string(pil_img)
                                page_text = (page_text or "") + "\n" + ocr_text
                                logger.debug("OCR used on page %d, extracted %d chars", page_num, len(ocr_text))
                            except Exception as e:
                                logger.debug("OCR failed on page %d: %s", page_num, e)
                        else:
                            logger.debug("No pytesseract available; skipping OCR")

                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error("Error reading PDF via pdfplumber: %s", e)
            return ""
        return text

    # -----------------------
    # Table extraction (camelot)
    # -----------------------
    def _extract_tables_from_pdf(self) -> List:
        try:
            tables = []
            try:
                if camelot:
                    tables = camelot.read_pdf(self.pdf_path, pages='all', flavor='lattice')
                    logger.info("Camelot lattice found %d tables", len(tables))
                else:
                    logger.debug("Camelot not available")
            except Exception as e:
                logger.debug("Camelot lattice failed: %s", e)

            if not tables and camelot:
                try:
                    tables = camelot.read_pdf(self.pdf_path, pages='all', flavor='stream')
                    logger.info("Camelot stream found %d tables", len(tables))
                except Exception as e:
                    logger.debug("Camelot stream failed: %s", e)

            for table in tables:
                try:
                    if hasattr(table, '_close_temp_files'):
                        table._close_temp_files()
                except Exception:
                    pass

            return tables or []
        except Exception as e:
            logger.error("Error extracting tables: %s", e)
            return []

    # -----------------------
    # Product info extraction
    # -----------------------
    def _extract_product_info(self) -> Dict:
        if model:
            return self._extract_product_info_with_ai()
        return self._extract_product_info_with_regex()

    def _extract_product_info_with_ai(self) -> Dict:
        try:
            prompt = f"""
Extract the following product information from this Process Validation Protocol:

Text:
{self.full_text[:4000]}

Return ONLY a JSON object with the fields:
{{"product_name": "...", "strength": "...", "dosage_form": "...", "batch_size": "...", "pack_size":"...", "manufacturing_site":"..."}}
"""
            response = model.generate_content(prompt)
            result_text = getattr(response, 'text', str(response)).strip()
            jstart = result_text.find('{')
            jend = result_text.rfind('}')
            if jstart != -1 and jend != -1 and jend > jstart:
                candidate = result_text[jstart:jend+1]
                try:
                    product_info = json.loads(candidate)
                    return product_info
                except Exception as e:
                    logger.error("AI returned non-JSON or parse failed: %s", e)
            logger.warning("AI did not return valid JSON, falling back to regex")
            return self._extract_product_info_with_regex()
        except Exception as e:
            logger.error("AI extraction failed: %s", e)
            return self._extract_product_info_with_regex()

    def _extract_product_info_with_regex(self) -> Dict:
        product_info = {
            'product_name': '',
            'strength': '',
            'dosage_form': '',
            'batch_size': '',
            'pack_size': '',
            'manufacturing_site': ''
        }

        lines = [l.strip() for l in self.full_text.splitlines() if l.strip()]

        # Product name
        for line in lines[:30]:
            if any(k in line.lower() for k in ['injection', 'tablet', 'capsule', 'syrup', 'suspension', 'vial']):
                product_info['product_name'] = line
                break

        # Strength
        m = re.search(r'(\d+(?:\.\d+)?)\s*(mg|g|ml|mcg|%|iu)(?:\s*/\s*(ml|g|tablet|capsule))?', self.full_text, re.I)
        if m:
            product_info['strength'] = m.group(0)

        # Batch size
        m = re.search(r'batch\s*size[:\s]*([\d, ]+\s*(?:vials|tablets|capsules|bottles|units|liters|litres)?)', self.full_text, re.I)
        if m:
            product_info['batch_size'] = m.group(1).strip()

        return product_info

    # -----------------------
    # Product type detection
    # -----------------------
    def _detect_product_type(self) -> str:
        text = (self.full_text or "").lower()
        injectable = sum(k in text for k in ['injection', 'injectable', 'vial', 'ampoule', 'aseptic'])
        tablet = sum(k in text for k in ['tablet', 'compression', 'granulation', 'coating'])
        capsule = sum(k in text for k in ['capsule', 'gelatin', 'shell'])
        oral = sum(k in text for k in ['syrup', 'suspension', 'solution', 'bottle filling'])
        scores = {'Injectable': injectable, 'Tablet': tablet, 'Capsule': capsule, 'Oral_Liquid': oral}
        return max(scores, key=scores.get)

    # -----------------------
    # Company info extraction
    # -----------------------
    def _extract_company_info(self) -> Dict:
        """Extract company information from the PVP document"""
        company_info = {
            'company_name': '',
            'company_address': '',
            'company_city': '',
            'company_state': '',
            'company_country': '',
            'company_pincode': ''
        }

        # Look for company name in first few lines
        lines = [l.strip() for l in self.full_text.splitlines() if l.strip()][:30]

        # common patterns for company name
        for line in lines:
            if any(k in line.upper() for k in ['LIMITED', 'LTD', 'PHARMACEUTICALS', 'PVT', 'INC', 'CORPORATION']):
                company_info['company_name'] = line
                break

        # Extract address with pincode (India-centric)
        address_match = re.search(
            r'([A-Z0-9\-,\s]+(?:AREA|ROAD|STREET|NAGAR|ZONE)[A-Z0-9\-,\s]+)[\s,]*([A-Z\s]+)[\s,\-]*(\d{6})',
            self.full_text,
            re.I
        )

        if address_match:
            company_info['company_address'] = address_match.group(1).strip()
            company_info['company_city'] = address_match.group(2).strip()
            company_info['company_pincode'] = address_match.group(3).strip()

        # Find state
        state_match = re.search(r'(DISTT?\.|DISTRICT|STATE)[:\s]*([A-Z\s]+)', self.full_text, re.I)
        if state_match:
            company_info['company_state'] = state_match.group(2).strip()

        # Find country
        if 'INDIA' in self.full_text.upper():
            company_info['company_country'] = 'INDIA'
        return company_info

    # -----------------------
    # Equipment extraction (FIXED)
    # -----------------------
    def _extract_equipment(self) -> List[Dict]:
        eq_from_tables = self._extract_equipment_from_tables()
        if eq_from_tables:
            return eq_from_tables

        # Fallback to text
        equipment = []
        pat = r'(?:equipment and machinery list|production equipment|equipment list)(.*?)(?:raw material|materials|quality control|$)'
        m = re.search(pat, self.full_text, re.I | re.S)
        if m:
            lines = [ln.strip() for ln in m.group(1).splitlines() if ln.strip()]
            for ln in lines:
                if len(ln) > 4 and not ln.lower().startswith('s.no') and not _is_obvious_heading(ln):
                    equipment.append({
                        'equipment_name': _clean_text_val(ln, max_len=190),
                        'equipment_id': '',
                        'location': '',
                        'calibration_status': ''
                    })
        return equipment

    def _extract_equipment_from_tables(self) -> List[Dict]:
        """Simpler, less aggressive filtering"""
        equipment = []
        seen = set()

        logger.info(f"=== EQUIPMENT EXTRACTION DEBUG ===")
        logger.info(f"Total tables to process: {len(self.tables_df)}")

        for idx, df in enumerate(self.tables_df):
            if df is None or df.shape[0] == 0:
                continue
            # Check if this looks like equipment table
            try:
                first_rows = ' '.join([str(x).lower() for x in df.iloc[:2].values.flatten()])

                logger.info(f"Table {idx}: shape={df.shape}, preview={first_rows[:100]}")

                # Look for equipment- related keywords
                equipment_keywords = ['equipment', 'machine', 'machinery', 'instrument', 'equipment id', 'equipment name', 'apparatus']
                is_equipment_table = any(kw in first_rows for kw in equipment_keywords) 

                if not is_equipment_table:
                    continue

                logger.info(f"Processing equipment table {idx} with {df.shape[0]} rows")

                header_row = 0
                if any(x in str(df.iloc[0, 0]).lower() for x in ['s.no', 'sr.no', 'serial']):
                    header_row = 1 if df.shape[0] > 1 else 0

                # Start processing from row after header

                start_row = header_row + 1
                for ridx in range(start_row, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()

                        non_empty = [x for x in row if str(x).strip() and str(x).strip().upper() != 'NAN']
                        if len(non_empty) < 2:
                            continue

                        eq_name = ''
                        eq_id = ''
                        # Try column 0 first
                        if len(row) > 0:
                            cell = str(row[0]).strip()
                    
                            # Skip serial numbers
                            if cell and not cell.isdigit() and len(cell) > 2:
                                eq_name = cell
                        if (not eq_name or eq_name.isdigit()) and len(row) > 1:
                            eq_name = str(row[1]).strip()

                        
                        # If column 0 looks like a number, try column 1
                        if len(row) > 2:
                            eq_id = str(row[2]).strip()
                        elif len(row) > 1 and eq_name != str(row[1]).strip():
                            eq_id = str(row[1]).strip()

                        # Clean and vadilate
                        eq_name_clean = _clean_text_val(eq_name, 190)
                        eq_id_clean = _clean_text_val(eq_id, 100)

                        # Skip invalid entries
                        if len(eq_name_clean) < 3:
                            continue

                        # skip obvious headers and junk
                        skip_terms = [
                            'equipment name', 'sr.no', 's.no', 'serial', 'equipment',
                            'format no', 'page no', 'instrument', 'machine', 
                            'effective date', 'supersedes', 'name of equipment'
                        ]

                        if any(term in eq_name_clean.lower() for term in skip_terms):
                            continue

                        # Skip it it's just "NAN" or similar
                        if eq_name_clean.upper() in ['NAN', 'NA', 'N/A', 'NONE', 'NULL', '-']:
                            continue

                        # Create unique key
                        key = (eq_name_clean.lower().strip(), eq_id_clean.lower().strip())

                        if key not in seen:
                            seen.add(key)
                            equipment.append({
                                'equipment_name': eq_name_clean,
                                'equipment_id': eq_id_clean if eq_id_clean and eq_id_clean.upper() not in ['N/A','NAN','NA'] else '',
                                'location': '',
                                'calibration_status': '',
                                'calibration_date': ''
                            })
                            logger.debug(f"Added equipment: {eq_name_clean}")
                    except Exception as e:
                        logger.debug(f"Error processing row {ridx}: {e}")
                        continue
            except Exception as e:
                logger.debug(f"Error processing table {idx}: {e}")
                continue
        logger.info(f"Total equipment extracted: {len(equipment)}")
        return equipment


    # -----------------------
    # Materials extraction
    # -----------------------
    def _extract_materials(self) -> List[Dict]:
        if model:
            try:
                return self._extract_materials_with_ai()
            except Exception as e:
                logger.debug("AI materials failed: %s", e)
        return self._extract_materials_with_regex()

    def _extract_materials_with_ai(self) -> List[Dict]:
        try:
            prompt = (
                "Extract all materials from this validation protocol. "
                "Categorize as API, Excipient, or Packaging.\n\n"
                "Text:\n"
                f"{self.full_text[:5000]}\n\n"
                "Return ONLY a JSON array like this:\n"
                '[\n'
                '  {\n'
                '    "material_type": "API",\n'
                '    "material_name": "Fluorouracil",\n'
                '    "specification": "USP",\n'
                '    "quantity": "500 mg"\n'
                '  }\n'
                ']\n\n'
                "Return only the JSON array, no other text."
            )
            response = model.generate_content(prompt)
            result_text = getattr(response, 'text', str(response)).strip()
            result_text = result_text.replace('```json', '').replace('```', '').strip()
            jstart = result_text.find('[')
            jend = result_text.rfind(']')
            if jstart != -1 and jend != -1 and jend > jstart:
                materials = json.loads(result_text[jstart:jend+1])
                logger.info(f"AI extracted {len(materials)} materials")
                return materials
        except Exception as e:
            logger.error(f"AI materials extraction failed: {e}")
        return self._extract_materials_with_regex()

    def _extract_materials_with_regex(self) -> List[Dict]:
        materials = []
        mats_from_tables = self._extract_materials_from_tables()
        if mats_from_tables:
            materials.extend(mats_from_tables)

        # Keyword heuristics
        text = (self.full_text or "").lower()
        apis = ['api', 'active ingredient', 'acetaminophen', 'paracetamol', 'ibuprofen', 'fluorouracil']
        excips = ['excipient', 'sodium hydroxide', 'water for injection', 'sodium chloride', 'preservative', 'tromethamine', 'disodium edetate']

        for k in apis:
            if k in text and not any(k in m.get('material_name', '').lower() for m in materials):
                materials.append({'material_type': 'API', 'material_name': k.title(), 'specification': '', 'quantity': ''})

        for k in excips:
            if k in text and not any(k in m.get('material_name', '').lower() for m in materials):
                materials.append({'material_type': 'Excipient', 'material_name': k.title(), 'specification': '', 'quantity': ''})

        return materials[:200]

    def _extract_materials_from_tables(self) -> List[Dict]:
        materials = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            if not headers:
                continue

            def find_idx(possibles):
                for p in possibles:
                    for i, h in enumerate(headers):
                        if p in h:
                            return i
                return None

            name_idx = find_idx(['material name', 'name', 'ingredient'])
            type_idx = find_idx(['material type', 'type'])
            spec_idx = find_idx(['specification', 'spec'])
            qty_idx = find_idx(['quantity', 'qty', 'amount'])

            if name_idx is not None:
                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    name = str(row[name_idx]).strip() if name_idx < len(row) else ''
                    if name and name.upper() != 'N/A' and len(name) > 2:
                        materials.append({
                            'material_type': str(row[type_idx]).strip() if type_idx is not None and type_idx < len(row) else '',
                            'material_name': _clean_text_val(name, max_len=190),
                            'specification': str(row[spec_idx]).strip() if spec_idx is not None and spec_idx < len(row) else '',
                            'quantity': str(row[qty_idx]).strip() if qty_idx is not None and qty_idx < len(row) else ''
                        })
        return materials

    # -----------------------
    # Batch details extraction
    # -----------------------
    def _extract_batch_details(self) -> List[Dict]:
        batches = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            if not headers:
                continue

            batch_keys = ['batch no', 'batch number', 'batch', 'lot', 'mfg date', 'manufacture', 'expiry', 'exp date']
            if any(any(k in h for h in headers) for k in batch_keys):
                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    row_dict = {}
                    for i, h in enumerate(headers):
                        row_dict[h] = str(row[i]).strip() if i < len(row) else ''
                    if row_dict:
                        batches.append(row_dict)
        return batches

    # -----------------------
    # NEW: Hold Time Study Extraction
    # -----------------------
    def _extract_hold_time_study(self) -> Dict:
        """Extract hold time study data (before/after filtration)"""
        hold_time = {
            'before_filtration': [],
            'after_filtration': []
        }

        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            # Look for hold time indicators
            if 'hold' in text or any(x in text for x in ['0 hour', '6 hour', '12 hour', '24 hour', '48 hour']):
                logger.info("Found hold time study table")

                for ridx in range(1, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()

                        # Extract data
                        time = str(row[0]).strip() if len(row) > 0 else ''
                        if not time or 'hour' not in time.lower():
                            continue

                        entry = {
                            'time': time,
                            'description': str(row[1]).strip() if len(row) > 1 else '',
                            'pH': str(row[2]).strip() if len(row) > 2 else '',
                            'assay': str(row[3]).strip() if len(row) > 3 else '',
                            'bioburden': str(row[4]).strip() if len(row) > 4 else '',
                            'sterility': str(row[5]).strip() if len(row) > 5 else ''
                        }

                        # Determine if before or after filtration
                        if 'before' in text or 'manufacturing tank' in text:
                            hold_time['before_filtration'].append(entry)
                        else:
                            hold_time['after_filtration'].append(entry)

                    except Exception as e:
                        logger.debug(f"Error processing hold time row: {e}")
                        continue

        return hold_time

    # -----------------------
    # NEW: Bioburden & BET Extraction
    # -----------------------
    def _extract_bioburden_bet(self) -> List[Dict]:
        """Extract bioburden and BET test results"""
        results = []

        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            if any(k in text for k in ['bioburden', 'bet', 'endotoxin', 'bacterial']):
                logger.info("Found bioburden/BET table")

                for ridx in range(1, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()

                        results.append({
                            'stage': _clean_text_val(str(row[0]).strip() if len(row) > 0 else '', 200),
                            'test': _clean_text_val(str(row[1]).strip() if len(row) > 1 else '', 100),
                            'specification': _clean_text_val(str(row[2]).strip() if len(row) > 2 else '', 200),
                            'result': _clean_text_val(str(row[3]).strip() if len(row) > 3 else '', 200)
                        })
                    except:
                        continue

        return results

    # -----------------------
    # NEW: Quality Control Final Product
    # -----------------------
    def _extract_qc_final_product(self) -> List[Dict]:
        """Extract final product quality control results"""
        qc_results = []

        # Search for QC section in text
        match = re.search(
            r'quality control.*?finish.*?product(.*?)(?:conclusion|deviation|summary|$)',
            self.full_text,
            re.I | re.S
        )

        if match:
            section = match.group(1)

            # Extract key tests
            tests_patterns = [
                ('Description', r'description[:\s]+(.*?)(?:\n|$)'),
                ('pH', r'ph[:\s]+([\d\.\-\s]+)'),
                ('Assay', r'assay[:\s]+([\d\.\-%\s]+)'),
                ('Sterility', r'sterility[:\s]+(.*?)(?:\n|$)'),
                ('BET', r'bet[:\s]+(.*?)(?:\n|$)'),
                ('Bacterial Endotoxins', r'bacterial endotoxins[:\s]+(.*?)(?:\n|$)'),
                ('Particulate Matter', r'particulate matter[:\s]+(.*?)(?:\n|$)'),
                ('Extractable Volume', r'extractable volume[:\s]+(.*?)(?:\n|$)')
            ]

            for test_name, pattern in tests_patterns:
                m = re.search(pattern, section, re.I)
                if m:
                    qc_results.append({
                        'test_name': test_name,
                        'result': _clean_text_val(m.group(1).strip(), 200)
                    })

        # Also check tables
        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            if 'quality' in text or 'finish' in text or 'specification' in text:
                for ridx in range(1, min(df.shape[0], 30)):
                    try:
                        row = df.iloc[ridx].tolist()
                        test_name = str(row[0]).strip() if len(row) > 0 else ''
                        result = str(row[1]).strip() if len(row) > 1 else ''

                        if test_name and result and len(test_name) > 2:
                            qc_results.append({
                                'test_name': _clean_text_val(test_name, 200),
                                'result': _clean_text_val(result, 200)
                            })
                    except:
                        continue

        return qc_results

    # -----------------------
    # NEW: Water Quality Extraction
    # -----------------------
    def _extract_water_quality(self) -> List[Dict]:
        """Extract water quality test results (purified water, WFI)"""
        water_tests = []

        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            if any(k in text for k in ['water', 'wfi', 'purified', 'conductivity']):
                logger.info("Found water quality table")

                for ridx in range(1, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()

                        water_tests.append({
                            'type': _clean_text_val(str(row[0]).strip() if len(row) > 0 else '', 100),
                            'test': _clean_text_val(str(row[1]).strip() if len(row) > 1 else '', 100),
                            'specification': _clean_text_val(str(row[2]).strip() if len(row) > 2 else '', 100),
                            'result': _clean_text_val(str(row[3]).strip() if len(row) > 3 else '', 100)
                        })
                    except:
                        continue

        return water_tests

    # -----------------------
    # NEW: Vial Sterility Extraction
    # -----------------------
    def _extract_vial_sterility(self) -> List[Dict]:
        """Extract vial sterility test results (initial/middle/end)"""
        sterility_tests = []

        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            if 'sterility' in text and any(k in text for k in ['initial', 'middle', 'end', 'vial']):
                logger.info("Found vial sterility table")

                for ridx in range(1, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()

                        sterility_tests.append({
                            'stage': _clean_text_val(str(row[0]).strip() if len(row) > 0 else '', 100),
                            'sample_size': _clean_text_val(str(row[1]).strip() if len(row) > 1 else '', 50),
                            'result': _clean_text_val(str(row[2]).strip() if len(row) > 2 else '', 100)
                        })
                    except:
                        continue

        return sterility_tests

    # -----------------------
    # NEW: Manufacturing Parameters
    # -----------------------
    def _extract_manufacturing_params(self) -> List[Dict]:
        """Extract manufacturing parameters (mixing time, stirring, etc.)"""
        params = []

        # Look for mixing/stirring parameters
        patterns = [
            (r'mixing.*?(\d+)\s*minutes', 'Mixing Time'),
            (r'stirring.*?(\d+)\s*rpm', 'Stirring Speed'),
            (r'temperature.*?([\d\.]+)\s*[°º]?\s*c', 'Temperature'),
            (r'pressure.*?([\d\.]+)\s*(?:kg/cm2|kg)', 'Pressure')
        ]

        for pattern, param_name in patterns:
            for m in re.finditer(pattern, self.full_text, re.I):
                params.append({
                    'parameter': param_name,
                    'value': m.group(1),
                    'context': m.group(0)
                })

        # Also check tables
        for df in self.tables_df:
            try:
                text = ' '.join([str(x).lower() for x in df.iloc[0].tolist()]) if df.shape[0] > 0 else ''
            except:
                continue

            if any(k in text for k in ['parameter', 'condition', 'mixing', 'stirring']):
                for ridx in range(1, df.shape[0]):
                    try:
                        row = df.iloc[ridx].tolist()
                        param = str(row[0]).strip() if len(row) > 0 else ''
                        value = str(row[1]).strip() if len(row) > 1 else ''

                        if param and value:
                            params.append({
                                'parameter': _clean_text_val(param, 200),
                                'value': _clean_text_val(value, 100),
                                'context': ''
                            })
                    except:
                        continue

        return params

    # -----------------------
    # Stages extraction
    # -----------------------
    def _extract_stages(self) -> List[Dict]:
        if model:
            try:
                return self._extract_stages_with_ai()
            except Exception as e:
                logger.debug("AI stages failed: %s", e)
        return self._extract_stages_with_regex()

    def _extract_stages_with_ai(self) -> List[Dict]:
        prompt = f"Extract manufacturing stages and return JSON array from the text:\n{self.full_text[:8000]}"
        try:
            response = model.generate_content(prompt)
            text = getattr(response, 'text', str(response)).strip()
            jstart = text.find('[')
            jend = text.rfind(']')
            if jstart != -1 and jend != -1 and jend > jstart:
                arr = json.loads(text[jstart:jend+1])
                return arr
        except Exception as e:
            logger.debug("AI stages parse failed: %s", e)
        return self._extract_stages_with_regex()

    def _extract_stages_with_regex(self) -> List[Dict]:
        stages = []
        major_process_headings = [
            ('Dispensing', r'dispensing'),
            ('Filtration', r'filtration'),
            ('Filling', r'filling'),
            ('Lyophilization', r'lyophilization'),
            ('Visual Inspection', r'visual inspection'),
            ('Sealing/Capping', r'seal|cap|capping'),
            ('Packaging', r'packaging'),
        ]

        for name, patt in major_process_headings:
            m = re.search(patt, self.full_text, re.I)
            if m:
                idx = m.start()
                context = self.full_text[max(0, idx - 200): idx + 400]
                stages.append({
                    'stage_number': len(stages) + 1,
                    'stage_name': name,
                    'equipment_used': '',
                    'parameters': context.strip()[:400],
                    'acceptance_criteria': ''
                })

        tbl_stages = self._extract_stages_from_tables()
        for s in tbl_stages:
            if not any(s['stage_name'].lower() in ex['stage_name'].lower() for ex in stages):
                stages.append(s)

        for i, st in enumerate(stages):
            st['stage_number'] = i + 1

        return stages

    def _extract_stages_from_tables(self) -> List[Dict]:
        stages = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            stage_keywords = ['stage', 'step', 'process', 'operation']
            if any(any(k in h for h in headers) for k in stage_keywords):
                def idx_of(poss):
                    for p in poss:
                        for i, h in enumerate(headers):
                            if p in h:
                                return i
                    return None

                name_idx = idx_of(['stage', 'step', 'process', 'activity'])
                eq_idx = idx_of(['equipment', 'machine'])
                param_idx = idx_of(['parameter', 'condition'])
                crit_idx = idx_of(['criteria', 'acceptance', 'limit'])

                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    name = str(row[name_idx]).strip() if name_idx is not None and name_idx < len(row) else ''
                    if name:
                        stages.append({
                            'stage_number': ridx,
                            'stage_name': _clean_text_val(name, max_len=200),
                            'equipment_used': _clean_text_val(str(row[eq_idx]).strip() if eq_idx is not None and eq_idx < len(row) else '', max_len=200),
                            'parameters': _clean_text_val(str(row[param_idx]).strip() if param_idx is not None and param_idx < len(row) else '', max_len=400),
                            'acceptance_criteria': _clean_text_val(str(row[crit_idx]).strip() if crit_idx is not None and crit_idx < len(row) else '', max_len=400)
                        })
        return stages

    # -----------------------
    # Test preparations & calculations
    # -----------------------
    def _extract_test_preparations(self) -> List[Dict]:
        preparations = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            if any('preparation' in h or 'absorbance' in h or 'area' in h for h in headers):
                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    preparations.append({
                        'test_name': _clean_text_val(str(row[0]).strip() if len(row) > 0 else '', max_len=200),
                        'preparation': _clean_text_val(str(row[1]).strip() if len(row) > 1 else '', max_len=400),
                        'area': _clean_text_val(str(row[2]).strip() if len(row) > 2 else '', max_len=100),
                        'absorbance': _clean_text_val(str(row[3]).strip() if len(row) > 3 else '', max_len=100)
                    })
        return preparations

    def _extract_calculation_sheet(self) -> List[Dict]:
        calculations = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            if any(k in ' '.join(headers) for k in ['mean', 'sd', 'rsd', 'calculation', 'result']):
                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    calculations.append({
                        'parameter': _clean_text_val(str(row[0]).strip() if len(row) > 0 else '', max_len=200),
                        'mean': _clean_text_val(str(row[1]).strip() if len(row) > 1 else '', max_len=100),
                        'sd': _clean_text_val(str(row[2]).strip() if len(row) > 2 else '', max_len=100),
                        'rsd': _clean_text_val(str(row[3]).strip() if len(row) > 3 else '', max_len=100),
                        'formula': _clean_text_val(str(row[4]).strip() if len(row) > 4 else '', max_len=400)
                    })
        return calculations

    # -----------------------
    # NEW: Statistics Calculation
    # -----------------------
    def _calculate_statistics_from_criteria(self, test_criteria: List[Dict]) -> Dict:
        """Calculate mean, SD, RSD for test results across batches"""
        stats = {}

        # Group by test name
        by_test = {}
        for test in test_criteria:
            test_name = test.get('test_name', '')
            acceptance = test.get('acceptance_criteria', '')

            # Try to extract numeric values
            numbers = re.findall(r'\d+\.?\d*', str(acceptance))
            if numbers:
                if test_name not in by_test:
                    by_test[test_name] = []
                by_test[test_name].extend([float(n) for n in numbers])

        # Calculate statistics
        for test_name, values in by_test.items():
            if len(values) > 0:
                mean = sum(values) / len(values)
                if len(values) > 1:
                    variance = sum((x - mean) ** 2 for x in values) / (len(values) - 1)
                    sd = variance ** 0.5
                    rsd = (sd / mean * 100) if mean != 0 else 0
                else:
                    sd = 0
                    rsd = 0

                stats[test_name] = {
                    'mean': round(mean, 2),
                    'sd': round(sd, 2),
                    'rsd': round(rsd, 2),
                    'count': len(values),
                    'min': round(min(values), 2),
                    'max': round(max(values), 2)
                }

        return stats

    # -----------------------
    # Observations, signatures, protocol summary
    # -----------------------
    def _extract_observations(self) -> str:
        m = re.search(r'(observations|remarks|deviations)[:\s]*\n?(.{0,800})', self.full_text, re.I)
        return m.group(2).strip() if m else ''

    def _extract_signatures(self) -> Dict:
        sigs = {}
        for role in ['performed by', 'checked by', 'approved by', 'authorized by']:
            m = re.search(fr'{role}[:\s]*([A-Za-z ,\.\-]+)', self.full_text, re.I)
            if m:
                sigs[role.replace(' ', '_')] = m.group(1).strip()

        m = re.search(r'(\d{2}/\d{2}/\d{4})', self.full_text)
        if m:
            sigs['date'] = m.group(1)

        return sigs

    def _extract_protocol_summary(self) -> str:
        start = re.search(r'(methodology|protocol|procedure)[:\s]*\n', self.full_text, re.I)
        if not start:
            return ''
        start_idx = start.start()
        end = re.search(r'(calculations|results|observations|conclusion|references)[:\s]*\n', self.full_text[start_idx:], re.I)
        if end:
            return self.full_text[start_idx:start_idx + end.start()].strip()
        return self.full_text[start_idx:].strip()

    # -----------------------
    # Test criteria extraction
    # -----------------------
    def _extract_test_criteria(self) -> List[Dict]:
        crits = []
        ct_from_tables = self._extract_test_criteria_from_tables()
        crits.extend(ct_from_tables)

        # Text heuristics
        patterns = [
            (r'pH[:\s]*([0-9\.]+)\s*(?:-|to|–)\s*([0-9\.]+)', 'pH'),
            (r'Assay[:\s]*([0-9\.]+)\s*(?:-|to|–)\s*([0-9\.]+)%', 'Assay'),
            (r'Volume[:\s]*([0-9\.]+)\s*ml', 'Volume')
        ]

        for patt, name in patterns:
            for m in re.finditer(patt, self.full_text, re.I):
                crits.append({
                    'test_id': f'test_{len(crits)+1}',
                    'test_name': name,
                    'acceptance_criteria': m.group(0)
                })

        return crits

    def _extract_test_criteria_from_tables(self) -> List[Dict]:
        criteria = []
        for df in self.tables_df:
            headers = [normalize_header(h) for h in df.iloc[0].tolist()] if df.shape[0] > 0 else []
            if not headers:
                continue

            if any(k in h for h in headers for k in ['test', 'parameter', 'specification', 'acceptance', 'limit']):
                def idx_of_poss(poss):
                    for p in poss:
                        for i, h in enumerate(headers):
                            if p in h:
                                return i
                    return None

                test_idx = idx_of_poss(['test', 'parameter', 'name'])
                crit_idx = idx_of_poss(['acceptance', 'criteria', 'limit', 'specification', 'range'])

                for ridx in range(1, df.shape[0]):
                    row = df.iloc[ridx].tolist()
                    if test_idx is not None and test_idx < len(row) and str(row[test_idx]).strip():
                        criteria.append({
                            'test_id': f'test_{len(criteria)+1}',
                            'test_name': _clean_text_val(str(row[test_idx]).strip(), max_len=200),
                            'acceptance_criteria': _clean_text_val(str(row[crit_idx]).strip() if crit_idx is not None and crit_idx < len(row) else '', max_len=400)
                        })
        return criteria

    # -----------------------
    # Template normalizer & docx generation
    # -----------------------
    def normalize_for_template(self, extracted: Dict) -> Dict:
        """Produce the JSON structure expected by templates"""
        out = {}
        out['meta'] = extracted.get('product_info', {})

        batches = extracted.get('batch_details', [])
        normalized_batches = []
        for b in batches:
            bn = b.get('batch no') or b.get('batch number') or b.get('batch') or b.get('lot') or ''
            if not bn:
                for v in b.values():
                    if str(v).strip():
                        bn = str(v).strip()
                        break
            normalized_batches.append({'batch_number': bn, **b})

        out['batches'] = normalized_batches

        # Test data
        test_ids = []
        data_rows = []
        for t in extracted.get('test_criteria', []):
            tid = t.get('test_id') or f"test_{len(test_ids)+1}"
            test_ids.append(tid)
            for b in normalized_batches:
                data_rows.append({
                    'test_id': tid,
                    'batch_number': b.get('batch_number'),
                    'test_result': t.get('acceptance_criteria', '')
                })

        out['test_ids'] = test_ids
        out['data'] = data_rows

        # Materials/Equipment tables
        mats = extracted.get('materials', [])
        if mats:
            headers = list(mats[0].keys())
            rows = [[m.get(h, '') for h in headers] for m in mats]
            out['materials_tables'] = [[headers] + rows]
        else:
            out['materials_tables'] = []

        eqs = extracted.get('equipment', [])
        if eqs:
            headers = list(eqs[0].keys())
            rows = [[e.get(h, '') for h in headers] for e in eqs]
            out['equipment_tables'] = [[headers] + rows]
        else:
            out['equipment_tables'] = []

        out['calculations'] = extracted.get('calculation_sheet', [])
        out['calculated'] = extracted.get('statistics', {})
        out['extracted'] = extracted
        out['generated_filepath'] = ''

        return out

    def generate_docx_from_extracted(self, extracted: Dict, out_path: str) -> Optional[str]:
        """Simple docx generator"""
        if Document is None:
            logger.warning("python-docx not installed, skipping DOCX generation")
            return None

        doc = Document()
        meta = extracted.get('product_info', {})
        title = meta.get('product_name') or "Process Validation Report"
        doc.add_heading(title, level=1)

        doc.add_paragraph(f"Protocol / Product info: {json.dumps(meta, ensure_ascii=False)}")

        # Batches
        doc.add_heading("Batch Details", level=2)
        batches = extracted.get('batch_details', [])
        if batches:
            keys = list(batches[0].keys())
            table = doc.add_table(rows=1, cols=len(keys))
            for i, k in enumerate(keys):
                table.rows[0].cells[i].text = str(k)
            for b in batches:
                r = table.add_row().cells
                for i, k in enumerate(keys):
                    r[i].text = str(b.get(k, ''))
        else:
            doc.add_paragraph("No batch details parsed.")

        # Materials
        doc.add_heading("Materials / Reagents", level=2)
        mats = extracted.get('materials', [])
        if mats:
            for m in mats:
                doc.add_paragraph(f"{m.get('material_type','')}: {m.get('material_name','')} - {m.get('quantity','')}")
        else:
            doc.add_paragraph("No materials parsed.")

        # Equipment
        doc.add_heading("Equipment", level=2)
        eqs = extracted.get('equipment', [])
        if eqs:
            for e in eqs:
                doc.add_paragraph(f"{e.get('equipment_name','')} (ID: {e.get('equipment_id','')})")
        else:
            doc.add_paragraph("No equipment parsed.")

        # Test Criteria
        doc.add_heading("Test Criteria", level=2)
        for t in extracted.get('test_criteria', []):
            doc.add_paragraph(f"{t.get('test_name','')}: {t.get('acceptance_criteria','')}")

        # Observations
        doc.add_heading("Observations", level=2)
        doc.add_paragraph(extracted.get('observations', '') or "-")

        # Signatures
        doc.add_heading("Signatures", level=2)
        doc.add_paragraph(json.dumps(extracted.get('signatures', {}), ensure_ascii=False))

        try:
            doc.save(out_path)
            logger.info("DOCX saved to %s", out_path)
            return out_path
        except Exception as e:
            logger.error("Failed to write DOCX: %s", e)
            return None

    # -----------------------
    # Helper: empty structure
    # -----------------------
    def _empty_result(self) -> Dict:
        return {
            'product_info': {},
            'product_type': 'Injectable',
            'equipment': [],
            'materials': [],
            'stages': [],
            'test_criteria': [],
            'batch_details': [],
            'hold_time_study': {'before_filtration': [], 'after_filtration': []},
            'bioburden_bet': [],
            'qc_final_product': [],
            'water_quality': [],
            'vial_sterility': [],
            'manufacturing_params': [],
            'statistics': {}
        }


def extract_from_pvp(pdf_path: str, output_dir: Optional[str] = None, tesseract_cmd: Optional[str] = None) -> Dict:
    """
    Convenience wrapper to run extraction and optionally create a DOCX.
    """
    extractor = EnhancedPVPExtractor(pdf_path, tesseract_cmd=tesseract_cmd)
    extracted = extractor.extract_all()
    template_payload = extractor.normalize_for_template(extracted)

    docx_path = None
    if output_dir:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        out_docx = os.path.join(output_dir, f"generated_pvr_{timestamp}.docx")
        docx_path = extractor.generate_docx_from_extracted(extracted, out_docx)
        template_payload['generated_filepath'] = docx_path or ''

    return {
        'extracted': extracted,
        'docx_path': docx_path,
        'template_payload': template_payload
    }


# -----------------------
# Example run / smoke test
# -----------------------
if __name__ == "__main__":
    sample_pdf = "/mnt/data/Fluorouracil Injection BP 50mg per ml Process Validation Report.pdf"
    outdir = "/tmp/pvp_output"

    print("Running extractor on:", sample_pdf)
    res = extract_from_pvp(sample_pdf, output_dir=outdir)
    if res and 'extracted' in res:
        print("Equipment count:", len(res['extracted'].get('equipment', [])))
        print("Product info:", res['extracted'].get('product_info'))
        print("Stages found:", len(res['extracted'].get('stages', [])))
    else:
        print("Extraction returned empty result or failed.")
