# Copyright (C) 2025 Soumyadeep Ghosh <soumyadeepghosh2004@zohomail.in>
# All Rights Reserved.

"""
AMV Protocol Generation System - Complete Backend with Protocol Structure
This system generates professional AMV protocols following ICH Q2(R1) guidelines
NO AI REQUIRED - Uses structured templates and mathematical calculations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from datetime import datetime, timedelta
import os
import json
import pandas as pd
import numpy as np
import random
import io


class AMVProtocolGenerator:
    # In the AMVProtocolGenerator class __init__ method, add proper JSON parsing:

    def __init__(self, form_data, verification_parameters=None, company_data=None):
        if not form_data:
            form_data = {
                'product_name': 'Test Product',
                'active_ingredient': 'Test Ingredient', 
            }

        self.form_data = form_data
        self.company_data = company_data or {}
        self.doc = Document()
        self.current_page = 1
        self.total_pages = 20
        self.setup_document_margins()
        
        # Parse JSON data from form
        self.selected_equipment = self._parse_json_data(form_data.get('selected_equipment_json'))
        self.selected_glass_materials = self._parse_json_data(form_data.get('selected_glass_materials_json'))
        self.selected_reagents = self._parse_json_data(form_data.get('selected_reagents_json'))
        self.selected_reference = self._parse_json_data(form_data.get('selected_reference_json'))
        
        # **FIXED: Handle validation parameters properly**
        self.verification_parameters = self._parse_validation_parameters(form_data, verification_parameters)
        
        print(f"🔍 FINAL VERIFICATION PARAMETERS: {self.verification_parameters}")


    def _parse_validation_parameters(self, form_data, verification_parameters):
        """Parse validation parameters from form data"""
        val_params = {}
        
        # First priority: Direct verification_parameters dict
        if verification_parameters and isinstance(verification_parameters, dict):
            print("🎯 Using direct verification_parameters dict")
            return verification_parameters
        
        # Second priority: val_params_json from form
        val_params_json = form_data.get('val_params_json')
        if val_params_json:
            try:
                selected_params = json.loads(val_params_json)
                print(f"📦 val_params_json found: {selected_params}")
                
                # Map form values to section numbers
                param_mapping = {
                    'system_suitability': '6.1',
                    'specificity': '6.2', 
                    'system_precision': '6.3',
                    'method_precision': '6.4',
                    'intermediate_precision': '6.5',
                    'linearity': '6.6',
                    'recovery': '6.7',
                    'robustness': '6.8',
                    'range': '6.9',
                    'lod_loq': '6.10',
                    'lod_loq_precision': '6.11'
                }
                
                for param in selected_params:
                    if param in param_mapping:
                        val_params[param_mapping[param]] = True
                        
            except Exception as e:
                print(f"❌ Error parsing val_params_json: {e}")
        
        # Third priority: Individual val_params from form (Flask ImmutableMultiDict)
        val_params_list = form_data.getlist('val_params') if hasattr(form_data, 'getlist') else []
        if val_params_list:
            print(f"📦 val_params list found: {val_params_list}")
            
            param_mapping = {
                'system_suitability': '6.1',
                'specificity': '6.2', 
                'system_precision': '6.3',
                'method_precision': '6.4',
                'intermediate_precision': '6.5',
                'linearity': '6.6',
                'recovery': '6.7',
                'robustness': '6.8',
                'range': '6.9',
                'lod_loq': '6.10',
                'lod_loq_precision': '6.11'
            }
            
            for param in val_params_list:
                if param in param_mapping:
                    val_params[param_mapping[param]] = True
        
        # If still no parameters selected, use sensible defaults based on instrument type
        if not val_params:
            test_method = form_data.get('test_method', 'HPLC').upper()
            print(f"⚠️ No parameters selected, using defaults for: {test_method}")
            
            if test_method in ['HPLC', 'UPLC', 'LC']:
                default_params = ['6.1', '6.2', '6.3', '6.4', '6.7']  # System Suitability, Specificity, System Precision, Method Precision, Accuracy
            elif test_method in ['TITRATION']:
                default_params = ['6.3', '6.4', '6.7']  # System Precision, Method Precision, Accuracy
            elif test_method in ['UV', 'SPECTROPHOTOMETRY']:
                default_params = ['6.2', '6.4', '6.7']  # Specificity, Method Precision, Accuracy
            else:
                default_params = ['6.1', '6.2', '6.4', '6.7']  # Generic defaults
            
            for param in default_params:
                val_params[param] = True
        
        print(f"✅ Final validation parameters: {val_params}")
        return val_params
    
    def _parse_json_data(self, data):
        """Parse JSON data safely from form inputs"""
        print(f"🔍 RAW JSON DATA: {repr(data)}")
        
        if not data or str(data).strip() in ['', 'null', 'None', '[]', '{}']:
            return []
        
        try:
            # Handle string JSON data
            if isinstance(data, str):
                # Clean the string data
                cleaned_data = data.strip()
                if not cleaned_data or cleaned_data in ['null', 'None', '[]', '{}']:
                    return []
                
                # Parse JSON
                parsed = json.loads(cleaned_data)
                print(f"✅ Parsed JSON: {parsed}")
                
                # Ensure we return a list
                if isinstance(parsed, list):
                    return parsed
                elif isinstance(parsed, dict):
                    return [parsed]
                else:
                    return []
            
            # Handle already parsed data
            elif isinstance(data, list):
                return data
            elif isinstance(data, dict):
                return [data]
            else:
                print(f"⚠️ Unknown data type: {type(data)}")
                return []
                
        except json.JSONDecodeError as e:
            print(f"❌ JSON decode error: {e}")
            print(f"❌ Problematic data: {repr(data)}")
            return []
        except Exception as e:
            print(f"❌ Error parsing JSON data: {e}")
            return []



    
    def setup_document_margins(self):
        """Set document margins safely"""
        try:
            sections = self.doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
        except Exception as e:
            print(f"Warning: Could not set margins: {e}")
    
    def add_header_section(self, page_number=None):
        """Add header with company info and product details - SAFE VERSION"""
        if page_number is not None:
            self.current_page = page_number
        
        try:
            # Company header table
            header_table = self.doc.add_table(rows=1, cols=2)
            header_table.autofit = False
            
            # Left cell - Company name and address
            left_cell = header_table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            
            company_name = self.form_data.get('company_name', '')
            company_address = self.form_data.get('company_location', '')
            
            company_run = left_para.add_run(f"{company_name}\n")
            company_run.bold = True
            company_run.font.size = Pt(10)
            
            address_run = left_para.add_run(f"{company_address}\n")
            address_run.font.size = Pt(9)
            
            # Protocol title
            title_run = left_para.add_run("ANALYTICAL METHOD VALIDATION PROTOCOL")
            title_run.bold = True
            title_run.font.size = Pt(10)
            
            # Right cell - Product info
            right_cell = header_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Product info table
            product_table = self.doc.add_table(rows=3, cols=2)
            product_table.style = 'Table Grid'
            
            # Product name
            product_table.rows[0].cells[0].text = "NAME OF PRODUCT"
            product_name = self.form_data.get('product_name', '')
            product_table.rows[0].cells[1].text = str(product_name)[:100]  # Limit length
            
            # Label claim
            product_table.rows[1].cells[0].text = "LABEL CLAIM"
            active_ingredient = self.form_data.get('active_ingredient', '')
            label_claim = self.form_data.get('label_claim', '')
            label_text = f"{active_ingredient}\t{label_claim}"
            product_table.rows[1].cells[1].text = str(label_text)[:200]  # Limit length
            
            # Protocol number and page
            protocol_no = self.form_data.get('protocol_number', '')
            product_table.rows[2].cells[0].text = f"PROTOCOL NO. {protocol_no}"
            product_table.rows[2].cells[1].text = f"PAGE {self.current_page} OF {self.total_pages}"
            
            self.doc.add_paragraph()
            
        except Exception as e:
            print(f"Error in header section: {e}")
            # Add basic header as fallback
            self.doc.add_paragraph("ANALYTICAL METHOD VALIDATION PROTOCOL")
    
    def add_page_break(self):
        """Add page break and increment page counter safely"""
        try:
            self.doc.add_page_break()
            self.current_page += 1
        except Exception as e:
            print(f"Error adding page break: {e}")
    
    def generate_protocol_cover(self):
        """Generate protocol cover page safely"""
        try:
            self.add_header_section(1)
            
            # Add title
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(
                f"\n\n\nANALYTICAL METHOD VALIDATION PROTOCOL FOR ASSAY\n"
                f"OF\n"
                f"{self.form_data.get('product_name', '')}"
            )
            title_run.bold = True
            title_run.font.size = Pt(14)
            
            self.add_page_break()
        except Exception as e:
            print(f"Error generating cover: {e}")
    
    def add_table_of_contents(self):
        """Add table of contents safely"""
        try:
            self.add_header_section(2)
            
            # Title
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("CONTENTS")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # TOC content
            toc_content = [
                ("1. Approval", 3),
                ("2. Overview", 4),
                ("2.1 Objective", 4),
                ("2.2 Scope", 4),
                ("2.3 Responsibility", 4),
                ("3. Details of Instruments/ Equipment", 4),
                ("3.1 List of Equipment and Instrument Used", 4),
                ("3.2 List of Glass or Other Materials", 4),
                ("3.3 List of Reagents and Prepared Solutions:", 5),
                ("3.4 Working Standard Details:", 5),
                ("4. References: ICH Q2 (R2)", 5),
                ("5. Methodology", 5),
                ("6. Validation Parameter", 7),
                ("6.1 System Precision", 8),
                ("6.2 Specificity", 9),
                ("6.3 Method Precision", 10),
                ("6.4 Intermediate precision (Ruggedness):", 12),
                ("6.5 Linearity and Range", 15),
                ("6.6 Accuracy/Recovery:", 17),
                ("6.7 Robustness", 19),
                ("7. Validation Report", 20)
            ]
            
            for content, page in toc_content:
                para = self.doc.add_paragraph()
                para.add_run(str(content))
                para.add_run(f"\t{page}")
                
                # Add indentation for sub-items
                if content.startswith(('2.', '3.', '6.')) and not content.startswith(('2.1', '3.1', '6.1')):
                    try:
                        para.paragraph_format.left_indent = Inches(0.3)
                    except:
                        pass
            
            self.add_page_break()
        except Exception as e:
            print(f"Error in table of contents: {e}")
    
    def add_approval_section(self):
        """Add approval section with signatures safely"""
        try:
            self.add_header_section(3)
            
            # Title
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("1. Approval")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # Description
            desc_para = self.doc.add_paragraph()
            product_name = self.form_data.get('product_name', '')
            method_name = self.form_data.get('test_method', '')
            desc_text = (
                f"This is a specific protocol for analytical method validation for Assay of "
                f"{product_name} by {method_name}.\n"
                f"This protocol has been approved by the following:"
            )
            desc_para.add_run(desc_text)
            
            # Approval table
            approval_table = self.doc.add_table(rows=4, cols=5)
            approval_table.style = 'Table Grid'
            
            # Headers
            headers = ['', 'Name', 'Department', 'Signature', 'Date']
            for idx, header in enumerate(headers):
                cell = approval_table.rows[0].cells[idx]
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
            
            # Get dates
            date_option = self.form_data.get('date_option', 'auto')
            if date_option == 'auto':
                protocol_date = datetime.now()
            else:
                try:
                    protocol_date = datetime.strptime(self.form_data.get('protocol_date', ''), '%Y-%m-%d')
                except:
                    protocol_date = datetime.now()
            
            date_str = protocol_date.strftime('%d/%m/%Y')
            
            # Approval data
            approval_data = [
                ('Prepared By', 
                 self.form_data.get('prepared_by_name', ''), 
                 'Analyst Q.C', 
                 '[Signature]', 
                 date_str),
                ('Checked By', 
                 self.form_data.get('reviewed_by_name', ''), 
                 'Asst. Manager Q.C', 
                 '[Signature]', 
                 date_str),
                ('Approved By', 
                 self.form_data.get('approved_by_name', ''), 
                 'Manager Q.C', 
                 '[Signature]', 
                 date_str)
            ]
            
            for idx, (role, name, dept, sig, date) in enumerate(approval_data, 1):
                if idx < len(approval_table.rows):  # Safety check
                    approval_table.rows[idx].cells[0].text = role
                    approval_table.rows[idx].cells[1].text = name
                    approval_table.rows[idx].cells[2].text = dept
                    approval_table.rows[idx].cells[3].text = sig
                    approval_table.rows[idx].cells[4].text = date
            
            self.add_page_break()
        except Exception as e:
            print(f"Error in approval section: {e}")
    
    def add_overview_section(self):
        """Add overview section safely"""
        try:
            self.add_header_section(4)
            
            # Main title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run("2. Overview")
            title_run.bold = True
            
            # Objective
            self.doc.add_paragraph("2.1 Objective")
            objective_text = (
                f"To establish the methodology for the analytical method validation for Assay of "
                f"{self.form_data.get('product_name', '')} "
                f"by {self.form_data.get('test_method', '')}."
            )
            self.doc.add_paragraph(objective_text)
            
            # Scope
            self.doc.add_paragraph("2.2 Scope")
            scope_text = (
                f"This Validation is applicable for the determination of Assay of "
                f"{self.form_data.get('product_name', '')} "
                f"by {self.form_data.get('test_method', '')}."
            )
            self.doc.add_paragraph(scope_text)
            
            # Responsibility
            self.doc.add_paragraph("2.3 Responsibility")
            responsibility_text = (
                "• Executive QC.\n"
                "- To prepare the analytical method validation protocol and to carry out the analytical work In accordance with this protocol.\n"
                "- To carry out all operations in accordance with GLP and relative Standard Operating Procedures (SOPs).\n"
                "- To record all observations.\n"
                "• Assistant Manager QC\n"
                "- To check the protocol and report.\n"
                "• Head Quality\n"
                "Head Quality will approve the Protocol and Report."
            )
            self.doc.add_paragraph(responsibility_text)
            
            # Equipment section title
            equip_title = self.doc.add_paragraph()
            equip_run = equip_title.add_run("3. Details of Instruments/ Equipment")
            equip_run.bold = True
            
            # Equipment list
            self.doc.add_paragraph("3.1 List of Equipment and Instrument Used:")
    
            if self.selected_equipment and len(self.selected_equipment) > 0:
                equip_text = "The following apparatus / equipment shall be used for validation studies:\n"
                for idx, equipment in enumerate(self.selected_equipment, 1):
                    letter = chr(96 + idx)  # Convert 1->'a', 2->'b', etc.
                    name = equipment.get('name', '')
                    code = equipment.get('code', '')
                    brand = equipment.get('brand', '')
                    
                    equip_text += f"{letter}) {name}"
                    if code:
                        equip_text += f" (Code: {code})"
                    if brand:
                        equip_text += f" - {brand}"
                    equip_text += "\n"
            else:
                # Default equipment if none selected
                equip_text = "The following apparatus / equipment shall be used for validation studies:\n" \
                            "a) Analytical Balance\n" \
                            "b) High Performance Liquid Chromatography\n" \
                            "c) Ultra-sonic Bath\n" \
                            "d) Vacuum Pump"
            
            self.doc.add_paragraph(equip_text)
            
            # Glass materials - USING ACTUAL SELECTED DATA
            self.doc.add_paragraph("3.2 List of Glass or Other Materials")

            if self.selected_glass_materials and len(self.selected_glass_materials) > 0:
                glass_text = "The following glass or other materials shall be used for Validation studies:\n"
                for idx, material in enumerate(self.selected_glass_materials, 1):
                    letter = chr(96 + idx)
                    name = material.get('name', '')
                    characteristics = material.get('characteristics', '')
                    glass_text += f"{letter}) {name}"
                    if characteristics:
                        glass_text += f" - {characteristics}"
                    glass_text += "\n"
            else:
                # Default materials if none selected
                glass_text = "The following glass or other materials shall be used for Validation studies:\n" \
                            "a) Beaker: 1000ml\n" \
                            "b) Glass Volumetric Flask: 50ml, 100ml\n" \
                            "c) Pipette: 2ml, 2.5ml, 5ml\n" \
                            "d) Graduated Cylinders: 500ml\n" \
                            "e) Glass jars: 1000ml"

            self.doc.add_paragraph(glass_text)

            self.add_page_break()
            
            # Reagents - USING ACTUAL SELECTED DATA
            self.doc.add_paragraph("3.3 List of Reagents and Prepared Solutions:")
            reagents_text = "The following reagents and chemicals shall be used for Validation studies:"
            self.doc.add_paragraph(reagents_text)

            try:
                reagents_table = self.doc.add_table(rows=1, cols=3)
                reagents_table.style = 'Table Grid'
                
                # Add headers
                header_cells = reagents_table.rows[0].cells
                header_cells[0].text = "Reagent Name"
                header_cells[1].text = "Batch Number"
                header_cells[2].text = "Expiry Date"
                
                # Use selected reagents or defaults
                if self.selected_reagents and len(self.selected_reagents) > 0:
                    reagents_data = self.selected_reagents
                else:
                    reagents_data = [
                        {
                            'name': 'Tetrabutylammonium hydroxide solution', 
                            'batch': 'TBH-001',
                            'expiry_date': '2025-10-31'
                        },
                        {
                            'name': 'Methanol', 
                            'batch': 'MTH-002',
                            'expiry_date': '2026-05-15'
                        }
                    ]
                
                for reagent in reagents_data:
                    row = reagents_table.add_row()
                    cells = row.cells
                    cells[0].text = reagent.get('name', '')
                    cells[1].text = reagent.get('batch', '')
                    
                    # Format expiry date
                    expiry_date = reagent.get('expiry_date', '')
                    if expiry_date:
                        try:
                            date_obj = datetime.strptime(expiry_date, '%Y-%m-%d')
                            cells[2].text = date_obj.strftime('%d/%m/%Y')
                        except:
                            cells[2].text = expiry_date
                    else:
                        cells[2].text = 'N/A'
                        
            except Exception as e:
                print(f"Error creating reagents table: {e}")
                self.doc.add_paragraph("Reagents: As per methodology")
            
            # Working standard - USING ACTUAL SELECTED REFERENCE
            self.doc.add_paragraph("3.4 Working Standard Details:")
            try:
                standard_table = self.doc.add_table(rows=1, cols=2)
                standard_table.style = 'Table Grid'
                
                # Add headers
                header_cells = standard_table.rows[0].cells
                header_cells[0].text = "Name of Working Standard"
                header_cells[1].text = "Potency"
                
                # Use selected reference or defaults
                if self.selected_reference and isinstance(self.selected_reference, dict):
                    standard_name = self.selected_reference.get('standard_name', self.form_data.get('active_ingredient', ''))
                    standard_potency = self.selected_reference.get('potency', self.form_data.get('standard_potency', ''))
                else:
                    standard_name = self.form_data.get('active_ingredient', '')
                    standard_potency = self.form_data.get('standard_potency', '')
                
                row = standard_table.add_row()
                cells = row.cells
                cells[0].text = standard_name
                cells[1].text = standard_potency
                    
            except Exception as e:
                print(f"Error creating standard table: {e}")
            
            # References
            self.doc.add_paragraph("4. References: ICH Q2 (R2)")
            
            # Methodology title
            method_title = self.doc.add_paragraph()
            method_run = method_title.add_run("5. Methodology")
            method_run.bold = True
            
            # Methodology content
            # Methodology content
            method_code = self.form_data.get('methodology_code', '')

            # Helper function for generating detailed methodologies
            def generate_detailed_methodology(form_data, method_type):
                """Generate detailed methodology based on method type"""
                
                if method_type == 'HPLC':
                    active = form_data.get('active_ingredient', '')
                    label_claim = form_data.get('label_claim', '')
                    weight_sample = form_data.get('weight_sample', '')
                    
                    return (
                        f"Standard solution: Accurately weigh and transfer about {form_data.get('weight_standard', '')} of {active} working standard "
                        f"into a {form_data.get('final_concentration_standard', '')} volumetric flask. Add about 50ml of diluent and sonicate to dissolve. "
                        f"Make up the volume with diluent and mix well.\n\n"
                        f"Sample solution: Weigh and powder {weight_sample} tablets. Transfer accurately weighed powder equivalent to {label_claim} of {active} "
                        f"into a {form_data.get('final_concentration_sample', '')} volumetric flask. Add about 50ml of diluent and sonicate for 15 minutes "
                        f"with intermittent shaking. Make up the volume with diluent and mix well. Filter the solution through 0.45µ PVDF syringe filter, "
                        f"discarding first few ml of filtrate.\n\n"
                        f"Procedure: Inject {form_data.get('injection_volume', '')} of blank, standard solution (six replicate injections) and sample solution "
                        f"(in duplicate) into the chromatograph. Record the chromatograms and measure the peak responses. "
                        f"Calculate the content of {active} per tablet."
                    )
                
                elif method_type == 'TITRATION':
                    active = form_data.get('active_ingredient', '')
                    label_claim = form_data.get('label_claim', '')
                    weight_sample = form_data.get('weight_sample', '')
                    molecular_weight = form_data.get('molecular_weight', '')
                    
                    return (
                        f"Weigh and powder {weight_sample} tablets. Add a quantity of the powder containing 1 g of {active} to 100ml of water, "
                        f"add 50ml of 1M hydrochloric acid VS and boil for 1 minute to remove the carbon dioxide. Cool and titrate the excess of acid "
                        f"with 1M sodium hydroxide VS using methyl orange solution as indicator. Each ml of 1M hydrochloric acid VS is equivalent to "
                        f"{molecular_weight}mg of {active}.\n\n"
                        f"Standard solution: Accurately weigh and transfer about {form_data.get('weight_standard', '')} of {active} working standard "
                        f"into a {form_data.get('final_concentration_standard', '')} volumetric flask. Add about 50ml of water and sonicate to dissolve. "
                        f"Make up the volume with water and mix well.\n\n"
                        f"Sample solution: Weigh and powder {weight_sample} tablets. Transfer accurately weighed powder equivalent to {label_claim} of {active} "
                        f"into a {form_data.get('final_concentration_sample', '')} volumetric flask. Add about 50ml of water and sonicate for 15 minutes "
                        f"with intermittent shaking. Make up the volume with water and mix well. Filter the solution through Whatman filter paper No. 41, "
                        f"discarding first few ml of filtrate.\n\n"
                        f"Procedure: Pipette appropriate volumes of standard and sample solutions and titrate as described above. "
                        f"Calculate the content of {active} per tablet."
                    )
                
                elif method_type == 'UV':
                    active = form_data.get('active_ingredient', '')
                    label_claim = form_data.get('label_claim', '')
                    weight_sample = form_data.get('weight_sample', '')
                    wavelength = form_data.get('wavelength', '')
                    
                    return (
                        f"Standard solution: Accurately weigh and transfer about {form_data.get('weight_standard', '')} of {active} working standard "
                        f"into a {form_data.get('final_concentration_standard', '')} volumetric flask. Add about 50ml of diluent and sonicate to dissolve. "
                        f"Make up the volume with diluent and mix well. Further dilute to get a concentration suitable for UV measurement.\n\n"
                        f"Sample solution: Weigh and powder {weight_sample} tablets. Transfer accurately weighed powder equivalent to {label_claim} of {active} "
                        f"into a {form_data.get('final_concentration_sample', '')} volumetric flask. Add about 50ml of diluent and sonicate for 15 minutes "
                        f"with intermittent shaking. Make up the volume with diluent and mix well. Filter the solution through Whatman filter paper No. 41, "
                        f"discarding first few ml of filtrate. Dilute appropriately to get a concentration suitable for UV measurement.\n\n"
                        f"Blank: Diluent\n\n"
                        f"Procedure: Measure the absorbance of standard and sample solutions at {wavelength} using the blank to set zero. "
                        f"Calculate the content of {active} per tablet using the formula:\n\n"
                        f"Assay (%) = (As/Ast) × (Wst/Ws) × (Ds/Dst) × (P/100) × (Avg. Wt./Label Claim) × 100\n\n"
                        f"Where:\n"
                        f"As = Absorbance of sample solution\n"
                        f"Ast = Absorbance of standard solution\n"
                        f"Wst = Weight of standard taken (mg)\n"
                        f"Ws = Weight of sample taken (mg)\n"
                        f"Ds = Dilution factor of sample\n"
                        f"Dst = Dilution factor of standard\n"
                        f"P = Potency of standard (%)\n"
                        f"Avg. Wt. = Average weight of tablets (mg)"
                    )
                
                return "Methodology details not available"

            # Generate method_text based on test method
            if self.form_data.get('test_method', '').upper() in ['HPLC', 'LC', 'UPLC']:
                detailed_methodology = generate_detailed_methodology(self.form_data, 'HPLC')
                
                method_text = (
                    f"Product\tMethodology Code\n"
                    f"{self.form_data.get('product_name', '')}\t{method_code}\n"
                    f"Chromatographic system:\n\n"
                    f"Mode\t: {self.form_data.get('mode', 'LC')}\n"
                    f"Detector\t: {self.form_data.get('detector', 'UV 280 nm')}\n"
                    f"Column\t: {self.form_data.get('column', '4.5 mm X 25 cm; 5 µm L1')}\n"
                    f"Injection volume\t: {self.form_data.get('injection_volume', '')}\n"
                    f"Autosampler\t: {self.form_data.get('autosampler_temp', '10°')}\n"
                    f"Column\t: {self.form_data.get('column_temp', '50°')}\n"
                    f"Flow rate\t: {self.form_data.get('flow_rate', '')}\n"
                    f"{self.form_data.get('solution_preparation', 'Solution A: Dissolve 2.6 ml of Tetrabutylammonium hydroxide solution (40% in water) and 2.8 gm of disodium hydrogen phosphate in 1000 ml of water. Adjust with phosphoric acid to a pH of 7.8.')} "
                    f"Mobile Phase: {self.form_data.get('mobile_phase', 'Methanol and Solution A (150:850)')}\n\n"
                    f"{detailed_methodology}"
                )

            # For Titration methods:
            elif self.form_data.get('test_method', '').upper() in ['TITRATION', 'TITRIMETRY']:
                detailed_methodology = generate_detailed_methodology(self.form_data, 'TITRATION')
                
                method_text = (
                    f"Product\tMethodology Code\n"
                    f"{self.form_data.get('product_name', '')}\t{method_code}\n"
                    f"Methodology By Titration:\n"
                    f"{detailed_methodology}"
                )

            # For UV/Spectrophotometry:
            elif self.form_data.get('test_method', '').upper() in ['UV', 'SPECTROPHOTOMETRY']:
                detailed_methodology = generate_detailed_methodology(self.form_data, 'UV')
                
                method_text = (
                    f"Product\tMethodology Code\n"
                    f"{self.form_data.get('product_name', '')}\t{method_code}\n"
                    f"UV Spectrophotometry:\n"
                    f"Wavelength\t: {self.form_data.get('wavelength', '')}\n\n"
                    f"{detailed_methodology}"
                )

            # Generic/Other methods:
            else:
                method_text = (
                    f"Product\tMethodology Code\n"
                    f"{self.form_data.get('product_name', '')}\t{method_code}\n"
                    f"{self.form_data.get('solution_preparation', 'Methodology details here')}"
                )

            # Use it in your doc:
            self.doc.add_paragraph(method_text)
            
            self.add_page_break()
            self.add_header_section(6)
            
            # Continue methodology
            method_cont_text = (
                "Procedure: Inject 10µl injection of each solution as given below:"
            )
            self.doc.add_paragraph(method_cont_text)
            
            # Injection sequence table - SAFE VERSION
            try:
                injection_table = self.doc.add_table(rows=4, cols=2)
                injection_table.style = 'Table Grid'
                injection_table.rows[0].cells[0].text = "Sample ID"
                injection_table.rows[0].cells[1].text = "No. of injection"
                
                injection_data = [
                    ("Blank", "01"),
                    ("Standard solution", "06"),
                    ("Sample solution", "03"),
                    ("Standard Solution_BKT", "01")
                ]
                
                for idx, (sample, injections) in enumerate(injection_data, 1):
                    if idx < len(injection_table.rows):  # Safety check
                        injection_table.rows[idx].cells[0].text = sample
                        injection_table.rows[idx].cells[1].text = injections
            except Exception as e:
                print(f"Error creating injection table: {e}")
            
            # System suitability
            self.doc.add_paragraph("System suitability:")
            try:
                suitability_table = self.doc.add_table(rows=3, cols=3)
                suitability_table.style = 'Table Grid'
                
                suitability_table.rows[0].cells[0].text = "Sr. No."
                suitability_table.rows[0].cells[1].text = "System suitability parameter"
                suitability_table.rows[0].cells[2].text = "Acceptance criteria"
                
                suitability_data = [
                    ("1", "Tailing factor", "NMT 2.0"),
                    ("2", "%RSD of area in the standard solution replicates.", "NMT 2.0%")
                ]
                
                for idx, (sr_no, param, criteria) in enumerate(suitability_data, 1):
                    if idx < len(suitability_table.rows):  # Safety check
                        suitability_table.rows[idx].cells[0].text = sr_no
                        suitability_table.rows[idx].cells[1].text = param
                        suitability_table.rows[idx].cells[2].text = criteria
            except Exception as e:
                print(f"Error creating suitability table: {e}")
            
            # Calculation
            self.doc.add_paragraph("Calculation:")
            # Calculation formulas based on test method
            if self.form_data.get('test_method', '').upper() in ['HPLC', 'LC', 'UPLC']:
                calc_text = (
                    f"Analysis Samples: Standard solution and Sample Solution\n\n"
                    f"Calculate the percentage of the labeled amount of {self.form_data.get('active_ingredient', '')} "
                    f"({self.form_data.get('molecular_formula', '')}) in the portion of "
                    f"{self.form_data.get('product_name', '')} taken:\n\n"
                    f"Result = (rU/rS) × (CS/CU) × (Mr1/Mr2) × 100\n\n"
                    f"Where:\n"
                    f"rU = peak response of {self.form_data.get('active_ingredient', '')} from the Sample solution\n"
                    f"rS = peak response of {self.form_data.get('active_ingredient', '')} from the Standard solution\n"
                    f"CS = concentration of {self.form_data.get('active_ingredient', '')} working standard in the Standard solution (mg/ml)\n"
                    f"CU = nominal concentration of {self.form_data.get('active_ingredient', '')} in the Sample solution (mg/mL)\n"
                    f"Mr1 = molecular weight of {self.form_data.get('active_ingredient', '')}, {self.form_data.get('molecular_weight', '')}\n"
                    f"Mr2 = molecular weight of {self.form_data.get('active_ingredient', '')} salt form, {self.form_data.get('molecular_weight_salt', self.form_data.get('molecular_weight', ''))}\n\n"
                    f"Acceptance Criteria: {self.form_data.get('specification_range', '')}"
                )

            elif self.form_data.get('test_method', '').upper() in ['TITRATION', 'TITRIMETRY']:
                calc_text = (
                    f"Analysis Samples: Standard solution and Sample Solution\n\n"
                    f"Calculate the percentage of the labeled amount of {self.form_data.get('active_ingredient', '')} "
                    f"in {self.form_data.get('product_name', '')}:\n\n"
                    f"Percentage Content = (V × M × F × {self.form_data.get('molecular_weight', '')} × 100) / (W × 1000)\n\n"
                    f"Where:\n"
                    f"V = Volume of titrant consumed (ml)\n"
                    f"M = Molarity of titrant\n"
                    f"F = Factor/Equivalence factor\n"
                    f"W = Weight of sample taken (mg)\n"
                    f"Molecular Weight = {self.form_data.get('molecular_weight', '')}\n\n"
                    f"Each ml of 1M titrant is equivalent to {self.form_data.get('molecular_weight', '')}mg of "
                    f"{self.form_data.get('active_ingredient', '')}\n\n"
                    f"Acceptance Criteria: {self.form_data.get('specification_range', '')}"
                )

            elif self.form_data.get('test_method', '').upper() in ['UV', 'SPECTROPHOTOMETRY']:
                calc_text = (
                    f"Analysis Samples: Standard solution and Sample Solution\n\n"
                    f"Calculate the percentage of the labeled amount of {self.form_data.get('active_ingredient', '')} "
                    f"in {self.form_data.get('product_name', '')}:\n\n"
                    f"Assay (%) = (As/Ast) × (Wst/Ws) × (Ds/Dst) × (P/100) × (Avg. Wt./Label Claim) × 100\n\n"
                    f"Where:\n"
                    f"As = Absorbance of sample solution at {self.form_data.get('wavelength', '')}\n"
                    f"Ast = Absorbance of standard solution at {self.form_data.get('wavelength', '')}\n"
                    f"Wst = Weight of standard taken = {self.form_data.get('weight_standard', '')}\n"
                    f"Ws = Weight of sample taken = {self.form_data.get('weight_sample', '')}\n"
                    f"Ds = Dilution factor of sample\n"
                    f"Dst = Dilution factor of standard\n"
                    f"P = Potency of standard = {self.form_data.get('potency', '')}\n"
                    f"Avg. Wt. = Average weight of tablets = {self.form_data.get('average_weight', '')}\n"
                    f"Label Claim = {self.form_data.get('label_claim', '')}\n\n"
                    f"Acceptance Criteria: {self.form_data.get('specification_range', '')}"
                )

            else:
                calc_text = (
                    f"Calculate the percentage of the labeled amount of {self.form_data.get('active_ingredient', '')} "
                    f"in {self.form_data.get('product_name', '')} as per the approved methodology.\n\n"
                    f"Acceptance Criteria: {self.form_data.get('specification_range', '')}"
                )

            self.doc.add_paragraph(calc_text)
            
            spec_range = self.form_data.get('specification_range', '')
            # Parse the range to extract lower and upper limits
            import re
            limits = re.findall(r'(\d+\.?\d*)\s*%', spec_range)
            if len(limits) >= 2:
                lower_limit = limits[0]
                upper_limit = limits[1]
            else:
                lower_limit = ""
                upper_limit = ""

            # Dynamic Limit paragraph
            limit_text = (
                f"Limit: It contains not less than {lower_limit} percent and not more than {upper_limit} percent "
                f"of the labeled amount of {self.form_data.get('active_ingredient', '')}."
            )
            self.doc.add_paragraph(limit_text)
            
            self.add_page_break()
        except Exception as e:
            print(f"Error in overview section: {e}")
    
    def add_validation_parameters_section(self):
        """Add validation parameters section that reflects user selections"""
        try:
            self.add_header_section(7)
            
            # Main title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run("6. Validation Parameter:")
            title_run.bold = True
            
            # Introduction
            intro_text = f"The {self.form_data.get('test_method','')} method is evaluated for following validation parameters:"
            self.doc.add_paragraph(intro_text)
            
            # Parameters table - USING ACTUAL SELECTIONS
            try:
                # Define all parameters with their mapping
                all_parameters = [
                    ("6.1", "System Suitability"),
                    ("6.2", "Specificity"), 
                    ("6.3", "System Precision"),
                    ("6.4", "Method Precision"),
                    ("6.5", "Intermediate Precision"),
                    ("6.6", "Linearity and Range"),
                    ("6.7", "Accuracy/Recovery"),
                    ("6.8", "Robustness"),
                    ("6.9", "Range"),
                    ("6.10", "LOD and LOQ"),
                    ("6.11", "LOD and LOQ Precision")
                ]
                
                # Filter to only include selected parameters
                selected_params = []
                for param_key, param_name in all_parameters:
                    if self.verification_parameters.get(param_key, False):
                        selected_params.append((str(len(selected_params) + 1), param_name))
                
                print(f"📋 Displaying {len(selected_params)} selected parameters in document")
                
                if selected_params:
                    # Create table with proper rows
                    params_table = self.doc.add_table(rows=len(selected_params) + 1, cols=2)
                    params_table.style = 'Table Grid'
                    
                    # Headers
                    params_table.rows[0].cells[0].text = "Sr.No."
                    params_table.rows[0].cells[1].text = "Validation Parameters"
                    
                    # Add selected parameters
                    for idx, (sr_no, param_name) in enumerate(selected_params, 1):
                        if idx < len(params_table.rows):
                            params_table.rows[idx].cells[0].text = sr_no
                            params_table.rows[idx].cells[1].text = param_name
                else:
                    # Fallback if no parameters selected (shouldn't happen with defaults)
                    self.doc.add_paragraph("No validation parameters selected.")
                    
            except Exception as e:
                print(f"Error creating parameters table: {e}")
                # Add fallback text
                fallback_text = "• System Suitability\n• Specificity\n• System Precision\n• Method Precision\n• Accuracy/Recovery"
                self.doc.add_paragraph(fallback_text)
            
            self.add_page_break()
        except Exception as e:
            print(f"Error in validation parameters section: {e}")
        
    def add_system_precision_section(self):
        """Add system precision section safely"""
        try:
                        self.doc.add_heading("6.1 System Precision:", level=2)
                        self.doc.add_paragraph("The system precision of method is demonstrated by injecting the Blank/diluent, and Standard solution. "
                                               "For preparation of diluent, blank solution (diluent), Standard solution and chromatographic conditions "
                                               "refer to section 6.0 i.e. analytical methods.")
                        self.doc.add_paragraph("Acceptance criteria:")
                        self.doc.add_paragraph("✓ System Suitability should meet the requirement.")
                        self.doc.add_paragraph("✓ The relative Standard deviation of the replicate injections obtained from six replicates of Standard solution should be not more than 2.0%.")
                        self.doc.add_paragraph("✓ Tailing factor obtain from Standard solution is NMT 2.0")
                        self.add_page_break()        
        except Exception as e:
            print(f"Error in system precision section: {e}")
    
    def add_specificity_section(self):
        """Add specificity section safely"""
        try:
            self.add_header_section(9)
            
            # Title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run("6.2 Specificity:")
            title_run.bold = True
            
            # Description
            desc_text = (
                "To ensure the interference from blank and placebo solution those is likely to be present at the peak due to in Standard and sample solution."
            )
            self.doc.add_paragraph(desc_text)
            
            # Solutions preparation
            active_ingredient = self.form_data.get('active_ingredient', '')
            solutions_text = (
                f"Standard solution: Taken 12.0 mg of {active_ingredient} RS/WS and transfer it into 100 ml volumetric flask add water to dissolve the content sonicate if necessary, make up volume with water up to 100ml.\n"
                f"Placebo Solution: Take placebo solution equivalent to sample (except API) in 100ml volumetric flask. Pipette 2.5ml of resulting solution and transfer to 100ml volumetric flask, make volume upto 100 ml with water.\n"
                f"Sample solution: Take 4 vials of sample and reconstitute with water shake well to dissolve and transfer the content to 100ml volumetric flask. Rinse the same vials 2 to 3 times with water and transfer to the same volumetric flask and make up volume 100ml with water. Further pipette 2.5ml of above solution and transfer to 50ml volumetric flask, make volume upto 50 ml with water."
            )
            self.doc.add_paragraph(solutions_text)
            
            # Procedure
            self.doc.add_paragraph("Procedure: Inject 10µl of the above solutions in HPLC and record the chromatogram and check peak purity.")
            
            # Sequence table - SAFE VERSION
            try:
                sequence_table = self.doc.add_table(rows=5, cols=2)
                sequence_table.style = 'Table Grid'
                sequence_table.rows[0].cells[0].text = "Sample"
                sequence_table.rows[0].cells[1].text = "Number of Injections"
                
                sequence_data = [
                    ("Blank Solution", "1"),
                    ("Standard Solution", "1"),
                    ("Sample Solution", "1"),
                    ("Standard Solution Bkt.", "1")
                ]
                
                for idx, (sample, injections) in enumerate(sequence_data, 1):
                    if idx < len(sequence_table.rows):  # Safety check
                        sequence_table.rows[idx].cells[0].text = sample
                        sequence_table.rows[idx].cells[1].text = injections
            except Exception as e:
                print(f"Error creating sequence table: {e}")
            
            # Acceptance criteria
            self.doc.add_paragraph("Acceptance Criteria")
            criteria_text = (
                "✓ System Suitability should meet the requirement.\n"
                "✓ No significant interfering peak should appear in the blank chromatogram at the retention time of the main peak. Peak Purity should pass."
            )
            self.doc.add_paragraph(criteria_text)
            
            self.add_page_break()
        except Exception as e:
            print(f"Error in specificity section: {e}")
    
    def generate_protocol(self, output_filename):
        """Generate the complete AMV protocol safely"""
        try:
            # Cover page
            self.generate_protocol_cover()
            
            # Table of contents
            self.add_table_of_contents()
            
            # Approval section
            self.add_approval_section()
            
            # Overview section
            self.add_overview_section()
            
            # Validation parameters
            self.add_validation_parameters_section()
            
            
        
            
            
            
            
            # Add remaining sections with safety checks
            section_mapping = {
                '6.1': self.add_system_precision_section,      # System Suitability
                '6.2': self.add_specificity_section,           # Specificity
                '6.3': self.add_system_precision_section,      # System Precision  
                '6.4': self.add_method_precision_section,      # Method Precision
                '6.5': self.add_intermediate_precision_section, # Intermediate Precision
                '6.6': self.add_linearity_range_section,       # Linearity and Range
                '6.7': self.add_accuracy_recovery_section,     # Accuracy/Recovery
                '6.8': self.add_robustness_section,            # Robustness
                '6.9': self.add_linearity_range_section,       # Range (reuse linearity)
                '6.10': self.add_lod_loq_section,              # LOD and LOQ
                '6.11': self.add_lod_loq_precision_section     # LOD and LOQ Precision
            }
            
            for param_key, section_method in section_mapping.items():
                if self.verification_parameters.get(param_key, False):
                    try:
                        print(f"📄 Generating section for parameter: {param_key}")
                        section_method()
                    except Exception as e:
                        print(f"Error in {section_method.__name__}: {e}")
            
            # Save document
            if isinstance(output_filename, io.BytesIO):
                self.doc.save(output_filename)
            else:
                self.doc.save(output_filename)
            
            return output_filename
            
        except Exception as e:
            print(f"Error generating protocol: {e}")
            # Return minimal document even if there's an error
            try:
                if isinstance(output_filename, io.BytesIO):
                    self.doc.save(output_filename)
                else:
                    self.doc.save(output_filename)
                return output_filename
            except:
                # Final fallback - create a basic document
                basic_doc = Document()
                basic_doc.add_heading('AMV Protocol - Basic Version', 0)
                basic_doc.add_paragraph(f'Product: {self.form_data.get("product_name", "Unknown")}')
                basic_doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                if isinstance(output_filename, io.BytesIO):
                    basic_doc.save(output_filename)
                else:
                    basic_doc.save(output_filename)
                return output_filename

    # Simplified versions of remaining methods to prevent index errors
    def add_method_precision_section(self):
        """Simplified method precision section"""
        try:
            self.add_header_section(10)
            self.doc.add_heading("6.3 Method Precision", level=2)
            self.doc.add_paragraph("Method precision will be evaluated using six sample preparations.")
            self.doc.add_paragraph("Acceptance criteria: %RSD ≤ 2.0%")
            self.add_page_break()
        except Exception as e:
            print(f"Error in method precision section: {e}")

    def add_intermediate_precision_section(self):
        """Simplified intermediate precision section"""
        try:
            self.add_header_section(12)
            self.doc.add_heading("6.4 Intermediate Precision", level=2)
            self.doc.add_paragraph("Intermediate precision will be evaluated by different analysts on different days.")
            self.doc.add_paragraph("Acceptance criteria: %RSD ≤ 2.0%")
            self.add_page_break()
        except Exception as e:
            print(f"Error in intermediate precision section: {e}")

    def add_linearity_range_section(self):
        """Simplified linearity and range section"""
        try:
            self.add_header_section(15)
            self.doc.add_heading("6.5 Linearity and Range", level=2)
            self.doc.add_paragraph("Linearity will be evaluated from 80% to 120% of target concentration.")
            self.doc.add_paragraph("Acceptance criteria: r ≥ 0.999")
            self.add_page_break()
        except Exception as e:
            print(f"Error in linearity section: {e}")

    def add_accuracy_recovery_section(self):
        """Simplified accuracy/recovery section"""
        try:
            self.add_header_section(17)
            self.doc.add_heading("6.6 Accuracy/Recovery", level=2)
            self.doc.add_paragraph("Recovery will be evaluated at 80%, 100%, and 120% levels.")
            self.doc.add_paragraph("Acceptance criteria: 98.0% - 102.0%")
            self.add_page_break()
        except Exception as e:
            print(f"Error in accuracy section: {e}")

    def add_robustness_section(self):
        """Simplified robustness section"""
        try:
            self.add_header_section(19)
            self.doc.add_heading("6.7 Robustness", level=2)
            self.doc.add_paragraph("Robustness will be evaluated for critical method parameters.")
            self.doc.add_paragraph("Acceptance criteria: No significant impact on results")
            self.add_page_break()
        except Exception as e:
            print(f"Error in robustness section: {e}")

    def add_validation_report_section(self):
        """Simplified validation report section"""
        try:
            self.add_header_section(20)
            self.doc.add_heading("7. Validation Report", level=1)
            self.doc.add_paragraph("A comprehensive validation report will be generated summarizing all results.")
        except Exception as e:
            print(f"Error in validation report section: {e}")
    def add_lod_loq_section(self):
        """Add LOD and LOQ section"""
        try:
            self.add_header_section(self.current_page)
            self.doc.add_heading("6.10 LOD and LOQ", level=2)
            self.doc.add_paragraph("Limit of Detection (LOD) and Limit of Quantitation (LOQ) will be determined.")
            self.doc.add_paragraph("Acceptance criteria: LOD: S/N ≥ 3, LOQ: S/N ≥ 10")
            self.add_page_break()
        except Exception as e:
            print(f"Error in LOD/LOQ section: {e}")

    def add_lod_loq_precision_section(self):
        """Add LOD and LOQ Precision section"""
        try:
            self.add_header_section(self.current_page)
            self.doc.add_heading("6.11 LOD and LOQ Precision", level=2)
            self.doc.add_paragraph("Precision at LOD and LOQ levels will be evaluated.")
            self.doc.add_paragraph("Acceptance criteria: %RSD ≤ 10.0% at LOQ")
            self.add_page_break()
        except Exception as e:
            print(f"Error in LOD/LOQ Precision section: {e}")


class AnalyticalMethodVerificationService:
    """Service to generate Analytical Method Verification protocols"""
    
    def __init__(self):
        """Initialize the service without requiring form_data"""
        pass

    def _create_fallback_protocol(self, protocol_data):
        """Create a basic fallback protocol when main generation fails"""
        try:
            print("🔄 Creating fallback protocol...")
            
            doc = Document()
            
            # Add basic header
            title = doc.add_heading('ANALYTICAL METHOD VERIFICATION PROTOCOL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add basic information
            doc.add_heading('Basic Information', level=1)
            
            basic_info = [
                ('Product Name', protocol_data.get('product_name', 'Unknown Product')),
                ('Active Ingredient', protocol_data.get('active_ingredient', 'Unknown')),
                ('Test Method', protocol_data.get('test_method', 'Unknown Method')),
                ('Protocol Number', protocol_data.get('protocol_number', 'Unknown')),
                ('Company', protocol_data.get('company_name', 'Unknown Company'))
            ]
            
            for label, value in basic_info:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))
            
            # Add error notice
            doc.add_heading('Notice', level=1)
            doc.add_paragraph(
                'This is a simplified protocol generated due to technical issues with the main generator. '
                'Please contact technical support for the full protocol.'
            )
            
            # Add generation info
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            print("✅ Fallback protocol created successfully")
            return buffer
            
        except Exception as e:
            print(f"❌ Error creating fallback protocol: {e}")
            # Ultimate fallback - return empty bytes
            return io.BytesIO()
    
    def generate_verification_protocol(self, method_info, protocol_data):
        """Generate protocol from form data - COMPLETELY FIXED VERSION"""
        try:
            print("🔍 DEBUG: generate_verification_protocol called")
            print(f"📦 protocol_data keys: {list(protocol_data.keys()) if protocol_data else 'None'}")
            
            # Ensure we have valid data
            if not protocol_data:
                protocol_data = {}
            if not method_info:
                method_info = {}

            # **CRITICAL FIX: Get validation parameters from the form**
            selected_val_params = protocol_data.get('val_params', [])
            print(f"🔍 RAW val_params: {selected_val_params} (type: {type(selected_val_params)})")

            # **FIXED: Handle different data types properly**
            if isinstance(selected_val_params, str):
                # Convert string to list
                if selected_val_params.strip() and selected_val_params.startswith('['):
                    try:
                        selected_val_params = json.loads(selected_val_params)
                    except:
                        selected_val_params = [p.strip() for p in selected_val_params.split(',') if p.strip()]
                else:
                    selected_val_params = [selected_val_params] if selected_val_params.strip() else []
            
            # Ensure it's a list
            if not isinstance(selected_val_params, list):
                selected_val_params = [selected_val_params] if selected_val_params else []

            print(f"🔍 PROCESSED val_params: {selected_val_params}")

            # **FIXED: Correct parameter mapping that matches HTML checkbox values**
            parameter_mapping = {
                # Frontend value -> Section number
                'system_suitability': '6.1',
                'specificity': '6.2', 
                'system_precision': '6.3',
                'method_precision': '6.4',
                'intermediate_precision': '6.5',
                'linearity': '6.6',
                'recovery': '6.7',
                'robustness': '6.8',
                'range': '6.9',
                'lod_loq': '6.10',
                'lod_loq_precision': '6.11'
            }
            
            # Display names for sections
            display_names = {
                '6.1': 'System Suitability',
                '6.2': 'Specificity',
                '6.3': 'System Precision', 
                '6.4': 'Method Precision',
                '6.5': 'Intermediate Precision',
                '6.6': 'Linearity and Range',
                '6.7': 'Accuracy/Recovery',
                '6.8': 'Robustness',
                '6.9': 'Range',
                '6.10': 'LOD and LOQ',
                '6.11': 'LOD and LOQ Precision'
            }

            # Initialize all parameters to False
            verification_parameters = {key: False for key in display_names.keys()}
            
            # **FIXED: Map frontend checkbox values to section numbers**
            for frontend_value in selected_val_params:
                frontend_value_str = str(frontend_value).strip()
                print(f"🔍 Processing frontend value: '{frontend_value_str}'")
                
                # Direct mapping from checkbox value to section number
                if frontend_value_str in parameter_mapping:
                    section_number = parameter_mapping[frontend_value_str]
                    verification_parameters[section_number] = True
                    print(f"✅ Mapped '{frontend_value_str}' -> '{section_number}'")
                
                # If it's already a section number
                elif frontend_value_str in display_names:
                    verification_parameters[frontend_value_str] = True
                    print(f"✅ Direct section number: '{frontend_value_str}'")
                
                else:
                    print(f"❌ No mapping found for: '{frontend_value_str}'")

            # Count selected parameters
            selected_count = sum(verification_parameters.values())
            print(f"📊 Total parameters selected from form: {selected_count}")
            print(f"🎯 Final verification_parameters: {verification_parameters}")

            # ONLY use defaults if NO parameters were selected in the form
            if selected_count == 0:
                # Get the ACTUAL test method from protocol_data
                actual_test_method = protocol_data.get('test_method', protocol_data.get('testMethod', 'HPLC'))
                print(f"⚠️ No parameters selected in form, using defaults for: {actual_test_method}")
                
                # Default parameters based on ACTUAL test method type
                if actual_test_method.upper() in ['HPLC', 'LC', 'UPLC']:
                    default_params = ["6.1", "6.2", "6.3", "6.4", "6.7"]
                    print("🎯 Using HPLC defaults")
                elif actual_test_method.upper() in ['TITRATION', 'TITRIMETRY']:
                    default_params = ["6.3", "6.4", "6.7"]
                    print("🎯 Using TITRATION defaults")
                elif actual_test_method.upper() in ['UV', 'SPECTROPHOTOMETRY']:
                    default_params = ["6.2", "6.4", "6.7"]
                    print("🎯 Using UV defaults")
                else:
                    default_params = ["6.1", "6.2", "6.4", "6.7"]
                    print("🎯 Using GENERIC defaults")
                
                # Update verification parameters with defaults
                for param in default_params:
                    if param in verification_parameters:
                        verification_parameters[param] = True
            else:
                print("✅ Using parameters selected in form (NOT defaults)")

            # **FIXED: Extract ALL form data with proper fallbacks**
            form_data_for_protocol = {
                # Basic Information
                'product_name': protocol_data.get('product_name', protocol_data.get('productName', 'Unknown Product')),
                'protocol_number': protocol_data.get('protocol_number', protocol_data.get('protocolNumber', 'AMV-PROTOCOL-001')),
                'test_method': protocol_data.get('test_method', protocol_data.get('testMethod', 'HPLC')),
                'active_ingredient': protocol_data.get('active_ingredient', protocol_data.get('activeIngredient', 'Unknown Ingredient')),
                'label_claim': protocol_data.get('label_claim', protocol_data.get('labelClaim', '')),
                'company_name': protocol_data.get('company_name', protocol_data.get('companyName', 'Unknown Company')),
                'company_location': protocol_data.get('company_location', protocol_data.get('companyLocation', '')),
                'specification_range': protocol_data.get('specification_range', protocol_data.get('specificationRange', '90.0% - 110.0%')),
                'methodology_code': protocol_data.get('methodology_code', protocol_data.get('methodologyCode', '')),
                'standard_potency': protocol_data.get('standard_potency', protocol_data.get('standardPotency', '')),
                
                # Method Parameters
                'weight_standard': protocol_data.get('weight_standard', ''),
                'weight_sample': protocol_data.get('weight_sample', ''),
                'final_concentration_standard': protocol_data.get('final_concentration_standard', ''),
                'final_concentration_sample': protocol_data.get('final_concentration_sample', ''),
                'potency': protocol_data.get('potency', ''),
                'average_weight': protocol_data.get('average_weight', ''),
                'weight_per_ml': protocol_data.get('weight_per_ml', ''),
                'wavelength': protocol_data.get('wavelength', ''),
                'molecular_weight': protocol_data.get('molecular_weight', ''),
                'molecular_formula': protocol_data.get('molecular_formula', ''),
                'reference_absorbance_standard': protocol_data.get('reference_absorbance_standard', ''),
                'reference_area_standard': protocol_data.get('reference_area_standard', ''),
                'flow_rate': protocol_data.get('flow_rate', ''),
                'injection_volume': protocol_data.get('injection_volume', ''),
                'reference_volume': protocol_data.get('reference_volume', ''),
                'weight_sample_gm': protocol_data.get('weight_sample_gm', ''),
                'standard_factor': protocol_data.get('standard_factor', ''),
                
                # Team Information
                'prepared_by_name': protocol_data.get('prepared_by_name', protocol_data.get('preparedByName', '')),
                'prepared_by_dept': protocol_data.get('prepared_by_dept', protocol_data.get('preparedByDept', 'Quality Control')),
                'reviewed_by_name': protocol_data.get('reviewed_by_name', protocol_data.get('reviewedByName', '')),
                'reviewed_by_dept': protocol_data.get('reviewed_by_dept', protocol_data.get('reviewedByDept', 'Quality Control')),
                'approved_by_name': protocol_data.get('approved_by_name', protocol_data.get('approvedByName', '')),
                'approved_by_dept': protocol_data.get('approved_by_dept', protocol_data.get('approvedByDept', 'Quality Assurance')),
                'authorized_by_name': protocol_data.get('authorized_by_name', protocol_data.get('authorizedByName', '')),
                'authorized_by_dept': protocol_data.get('authorized_by_dept', protocol_data.get('authorizedByDept', 'Quality Assurance')),
                
                # Additional parameters
                'mode': protocol_data.get('mode', 'LC'),
                'detector': protocol_data.get('detector', 'UV 280 nm'),
                'column': protocol_data.get('column', '4.5 mm X 25 cm; 5 µm L1'),
                'autosampler_temp': protocol_data.get('autosampler_temp', '10°'),
                'column_temp': protocol_data.get('column_temp', '50°'),
                'mobile_phase': protocol_data.get('mobile_phase', 'Methanol and Solution A (150:850)'),
                'solution_preparation': protocol_data.get('solution_preparation', 'Solution A: Dissolve 2.6 ml of Tetrabutylammonium hydroxide solution (40% in water) and 2.8 gm of disodium hydrogen phosphate in 1000 ml of water. Adjust with phosphoric acid to a pH of 7.8.'),
                
                'date_option': 'auto'
            }

            # **FIXED: Handle JSON data properly**
            # Extract JSON data from form
            selected_equipment_json = protocol_data.get('selected_equipment_json', '[]')
            selected_glass_materials_json = protocol_data.get('selected_glass_materials_json', '[]')
            selected_reagents_json = protocol_data.get('selected_reagents_json', '[]')
            selected_reference_json = protocol_data.get('selected_reference_json', '{}')

            print(f"📦 Equipment JSON: {selected_equipment_json}")
            print(f"📦 Glass JSON: {selected_glass_materials_json}")
            print(f"📦 Reagents JSON: {selected_reagents_json}")
            print(f"📦 Reference JSON: {selected_reference_json}")

            # Add JSON data to form_data_for_protocol
            form_data_for_protocol.update({
                'selected_equipment_json': selected_equipment_json,
                'selected_glass_materials_json': selected_glass_materials_json,
                'selected_reagents_json': selected_reagents_json,
                'selected_reference_json': selected_reference_json,
            })

            print(f"🎯 Final verification_parameters: {verification_parameters}")
            print(f"📝 Test Method: {form_data_for_protocol.get('test_method')}")
            print(f"🏭 Product: {form_data_for_protocol.get('product_name')}")
            print(f"🔬 Protocol No: {form_data_for_protocol.get('protocol_number')}")
            print(f"👤 Prepared By: {form_data_for_protocol.get('prepared_by_name')}")

            # Generate document using the main protocol generator with the CORRECT form data
            generator = AMVProtocolGenerator(form_data_for_protocol, verification_parameters)
            buffer = io.BytesIO()
            generator.generate_protocol(buffer)
            buffer.seek(0)
            
            print("✅ Protocol generated successfully!")
            return buffer
            
        except Exception as e:
            print(f"❌ Error in generate_verification_protocol: {e}")
            import traceback
            traceback.print_exc()
            # Return a basic document as fallback
            return self._create_fallback_protocol(protocol_data)
        
    def generate_protocol_from_files(self, excel_path, pdf_path):
        """Generate protocol from Excel and PDF files"""
        try:
            # Create a basic protocol document
            doc = Document()
            doc.add_heading('AMV Protocol Generated from Files', 0)
            doc.add_paragraph(f'Excel file: {excel_path}')
            doc.add_paragraph(f'PDF file: {pdf_path}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer, "file_based"
        except Exception as e:
            print(f"Error generating protocol from files: {e}")
            return None, f"Error: {str(e)}"


# Create the service instance
analytical_method_verification_service = AnalyticalMethodVerificationService()


# Utility function for protocol generation
def generate_protocol_from_form(form_data):
    """Generate AMV protocol from form data - SAFE VERSION"""
    try:
        # Set default values if not provided
        defaults = {
            'company_name': 'KWALITY PHARMACEUTICALS. LTD.',
            'company_address': '1-A, INDUSTRIAL AREA, RAJA KA BAGH TEHSIL NURPUR, KANGRA-176201 (INDIA)',
            'product_name': 'LEUCOVORIN CALCIUM FOR INJECTION 50MG/VIAL',
            'active_ingredient': 'LEUCOVORIN CALCIUM',
            'label_claim': '50MG',
            'protocol_number': 'AMV/P/0154',
            'methodology_code': 'KPL/STP/IN/108-00',
            'standard_potency': '90.8%',
            'prepared_by': 'Sachin Kumar',
            'checked_by': 'Naresh',
            'approved_by': 'Ajay Bhatia',
            'date_option': 'auto'
        }
        
        # Merge form data with defaults
        for key, value in defaults.items():
            if key not in form_data or not form_data[key]:
                form_data[key] = value
        
        # Generate protocol
        generator = AMVProtocolGenerator(form_data)
        output_file = f"AMV_Protocol_{form_data['protocol_number']}.docx"
        result_file = generator.generate_protocol(output_file)
        
        return {
            'success': True,
            'file_path': result_file,
            'message': 'AMV Protocol generated successfully'
        }
        
    except Exception as e:
        print(f"Error in generate_protocol_from_form: {e}")
        return {
            'success': False,
            'error': str(e),
            'message': f'Error generating AMV Protocol: {str(e)}'
        }


if __name__ == "__main__":
    # Test the protocol generator
    test_data = {
        'product_name': 'LEUCOVORIN CALCIUM FOR INJECTION 50MG/VIAL',
        'active_ingredient': 'LEUCOVORIN CALCIUM',
        'label_claim': '50MG',
        'protocol_number': 'AMV/P/0154',
        'company_name': 'KWALITY PHARMACEUTICALS. LTD.',
        'company_address': '1-A, INDUSTRIAL AREA, RAJA KA BAGH TEHSIL NURPUR, KANGRA-176201 (INDIA)',
        'methodology_code': 'KPL/STP/IN/108-00',
        'standard_potency': '90.8%',
        'prepared_by': 'Sachin Kumar',
        'checked_by': 'Naresh',
        'approved_by': 'Ajay Bhatia'
    }
    
    result = generate_protocol_from_form(test_data)
    print(result['message'])